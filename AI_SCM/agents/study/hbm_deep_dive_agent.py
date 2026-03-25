"""
HBM Deep Dive Study Agent
대상: 반도체 마케팅 종사자 (비즈니스 배경 ○, 기술 기초 보완 필요)

학습 목표:
  - HBM이 왜 AI에 필수인지 기술적 원리를 마케팅 언어로 이해
  - HBM 시장 구조·경쟁·ASP 동학을 정량적으로 파악
  - 투자 시그널과 카탈리스트를 직접 도출할 수 있게 됨

출력: outputs/reports/study_hbm_deep_dive_YYYYMMDD.docx
"""
import os, sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
import config

OUTPUT_DIR = Path(__file__).parent.parent.parent / "outputs" / "reports"

# ──────────────────────────────────────────────────────────────
# 참고자료 (HBM 전문)
# ──────────────────────────────────────────────────────────────
REFS = {
    "SKHynix_HBM_Blog": (
        "SK Hynix Newsroom — HBM 기술 블로그",
        "https://news.skhynix.com/hbm/",
        "HBM 구조·TSV·세대별 기술 상세 설명 (한국어 포함)", "2023-2025"
    ),
    "SKHynix_IR_2024": (
        "SK Hynix IR — FY2024 실적발표",
        "https://www.skhynix.com/eng/ir/financialHighlights.do",
        "HBM 출하량, ASP, 가동률 공시. '수요가 공급을 초과' 반복 언급", "2024-2025"
    ),
    "Micron_HBM3E": (
        "Micron Technology — HBM3E 제품 페이지",
        "https://www.micron.com/products/high-bandwidth-memory",
        "Micron HBM3e 8-Hi/12-Hi 스펙, NVIDIA 공급 계획", "2024"
    ),
    "Omdia_HBM_2024": (
        "Omdia — HBM Market Monitor 2024",
        "https://omdia.tech.informa.com/OM026117/HBM-Market-Monitor",
        "HBM 시장규모 2019-2027E, 공급사 점유율, 세대별 수요 분석", "2024"
    ),
    "TrendForce_HBM": (
        "TrendForce — HBM Supply & Demand Tracker",
        "https://www.trendforce.com/research/dram.html",
        "분기별 HBM 공급·수요·가격 트렌드", "2024-2025"
    ),
    "NVIDIA_GTC_B200": (
        "NVIDIA GTC 2024 — B200/GB200 NVL72 발표",
        "https://www.nvidia.com/en-us/events/gtc/",
        "B200 HBM3e 192GB 탑재; GB200 NVL72 랙 스펙 공개", "2024"
    ),
    "TSMC_CoWoS_2024": (
        "TSMC — CoWoS Advanced Packaging",
        "https://pr.tsmc.com/english/news/3141",
        "CoWoS 구조, HBM과의 패키징 방식, 캐파 확대 계획", "2024"
    ),
    "SemiAnalysis_HBM": (
        "SemiAnalysis — HBM Supply Chain Deep Dive",
        "https://www.semianalysis.com/p/hbm-supply-chain-deep-dive",
        "HBM 제조 공정, CoWoS 의존성, SK Hynix 경쟁우위 분석", "2024"
    ),
    "Patterson_Memory_Wall": (
        "Patterson & Hennessy — Computer Architecture (Memory Wall)",
        "https://www.elsevier.com/books/computer-organization-and-design/patterson/978-0-12-820331-6",
        "메모리 벽(Memory Wall) 개념: CPU/GPU 연산 속도 vs 메모리 대역폭 격차", "2020"
    ),
    "Kaplan_Scaling": (
        "Kaplan et al. — Scaling Laws for Neural Language Models (ArXiv)",
        "https://arxiv.org/abs/2001.08361",
        "모델 크기 증가 → 성능 멱함수 향상 → HBM 수요 구조적 증가 근거", "2020"
    ),
    "Goldman_AI_Infra": (
        "Goldman Sachs — AI Infrastructure 2024",
        "https://www.goldmansachs.com/insights/articles/AI-poised-to-drive-160-increase-in-power-demand",
        "AI 공급망 투자 분석; HBM을 핵심 병목으로 지목", "2024"
    ),
    "Bernstein_HBM": (
        "Bernstein Research — HBM Market Analysis",
        "https://www.bernsteinresearch.com/",
        "HBM ASP 프리미엄, SK Hynix 경쟁 우위 정량 분석", "2024"
    ),
    "MLCommons_Bench": (
        "MLCommons MLPerf Inference v4.0",
        "https://mlcommons.org/benchmarks/inference-datacenter/",
        "H100/B200 추론 성능 벤치마크; 메모리 대역폭이 병목임을 실측", "2024"
    ),
}


# ──────────────────────────────────────────────────────────────
# 헬퍼 함수
# ──────────────────────────────────────────────────────────────
def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def _add_table(doc, headers, rows):
    if not rows:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        c._tc.get_or_add_tcPr().append(shd)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
            if tr.cells[ci].paragraphs[0].runs:
                tr.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

def _box(doc, label, color_fill, text):
    """색상 박스 (개념 정의, 마케터 관점, 기술 원리 구분)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_fill)
    p._p.get_or_add_pPr().append(shd)
    r1 = p.add_run(f"[{label}] ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.italic = True

def _concept(doc, text):
    """파란 박스: 기술 개념 설명."""
    _box(doc, "개념", "E8F0FE", text)

def _marketing(doc, text):
    """초록 박스: 마케터 관점 해석."""
    _box(doc, "마케터 관점", "E8F5E9", text)

def _quant(doc, text):
    """노란 박스: 정량 계산 예제."""
    _box(doc, "계산 예제", "FFF9C4", text)

def _risk(doc, text):
    """주황 박스: 리스크/주의사항."""
    _box(doc, "주의", "FFF3E0", text)

def _add_refs(doc, keys):
    refs = [REFS[k] for k in keys if k in REFS]
    if not refs:
        return
    doc.add_paragraph("참고자료", style="Heading 3")
    rows = [(n, url, desc, yr) for n, url, desc, yr in refs]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["출처", "URL", "설명", "연도"]):
        c = tbl.rows[0].cells[i]
        c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for ri, (name, url, desc, yr) in enumerate(rows):
        tr = tbl.rows[ri + 1]
        tr.cells[0].text = name
        p = tr.cells[1].paragraphs[0]
        run = p.add_run(url)
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.size = Pt(7)
        tr.cells[2].text = desc
        tr.cells[3].text = yr
        for ci in [0, 2, 3]:
            if tr.cells[ci].paragraphs[0].runs:
                tr.cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
    doc.add_paragraph()

def _quiz_section(doc, questions):
    """복습 문제 섹션."""
    _add_heading(doc, "복습 문제 (능동적 학습)", 2)
    doc.add_paragraph(
        "아래 질문에 먼저 스스로 답해보고, 힌트를 확인하세요. "
        "답을 '설명할 수 있으면' 이해한 것입니다 (페인만 기법)."
    )
    for i, (q, hint, answer_type) in enumerate(questions, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"Q{i}. {q}")
        r.bold = True
        r.font.size = Pt(10)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"  힌트: {hint}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(80, 80, 80)
        r2.italic = True
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"  유형: {answer_type}")
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(120, 120, 120)
        doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# 메인 빌드 함수
# ──────────────────────────────────────────────────────────────
def build():
    if not DOCX_AVAILABLE:
        print("  [HBM Study] python-docx 미설치")
        return None

    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_heading("HBM 심층 분석", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("High Bandwidth Memory — 기술 원리부터 투자 시그널까지")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = doc.add_paragraph(
        f"대상: 반도체 마케팅 종사자  |  기준일: {config.AS_OF_DATE}  |  "
        f"난이도: 기술 기초 → 중급"
    )
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s2.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    desc = doc.add_paragraph(
        "\n본 문서는 HBM(High Bandwidth Memory)을 반도체 마케팅 관점에서 이해하기 위한 "
        "심층 학습 자료입니다. 기술 원리를 비즈니스 언어로 풀어 설명하며, "
        "실제 계산 예제와 복습 문제로 실전 적용력을 높입니다."
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in desc.runs:
        run.font.size = Pt(9)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1장. 왜 HBM인가 — 메모리 병목의 본질
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "1장. 왜 HBM인가 — 메모리 병목의 본질")

    doc.add_paragraph(
        "HBM을 이해하려면 먼저 'AI가 왜 일반 메모리(DDR)로 안 되는가'를 알아야 합니다. "
        "답은 단 하나: 속도(대역폭, Bandwidth)입니다."
    )

    _add_heading(doc, "1.1 메모리 대역폭이란 무엇인가", 2)
    _concept(doc,
        "대역폭(Bandwidth)은 메모리가 CPU/GPU로 데이터를 전달하는 속도입니다. "
        "단위는 GB/s(초당 기가바이트). 수치가 클수록 같은 시간에 더 많은 데이터를 처리합니다.")
    _marketing(doc,
        "마케팅 비유: 대역폭 = 도로의 차선 수. "
        "DDR5 = 2차선 국도(89 GB/s), HBM3e = 60차선 고속도로(1,229 GB/s). "
        "화물(데이터)이 같아도 고속도로가 14배 빨리 나릅니다.")

    bw_rows = [
        ["DDR4 (일반 PC RAM)", "~51 GB/s",   "1배 (기준)", "$3-4/GB",  "PC, 서버 범용"],
        ["DDR5 (최신 서버)",   "~89 GB/s",   "1.7배",      "$5-7/GB",  "최신 서버 CPU"],
        ["GDDR6X (게임 GPU)", "~512 GB/s",  "10배",       "$6-8/GB",  "NVIDIA 게임 GPU"],
        ["HBM2e",              "~461 GB/s",  "9배",        "$8/GB",    "A100 GPU"],
        ["HBM3",               "~819 GB/s",  "16배",       "$14/GB",   "H100 GPU"],
        ["HBM3e",              "~1,229 GB/s","24배",       "$18/GB",   "H200, B200 GPU"],
        ["HBM4 (2025~)",       "~2,000 GB/s","39배",       "~$24/GB 추정", "Blackwell Ultra"],
    ]
    _add_table(doc, ["메모리 종류", "대역폭", "DDR4 대비", "가격/GB", "주요 용도"], bw_rows)

    _marketing(doc,
        "가격이 DDR4 대비 4-6배 비쌈에도 AI GPU에 HBM을 쓰는 이유: "
        "대역폭 24배 차이가 AI 추론 성능을 직접 결정하기 때문입니다. "
        "GPT-4 같은 모델은 초당 수조 개의 파라미터를 메모리에서 읽어야 하는데, "
        "DDR5로는 물리적으로 불가능합니다.")

    _add_heading(doc, "1.2 AI가 왜 메모리 대역폭에 민감한가", 2)
    doc.add_paragraph(
        "AI 모델(LLM)의 추론 과정에서 가장 많은 시간을 소비하는 작업은 "
        "'가중치(Weight) 읽기'입니다."
    )
    _concept(doc,
        "가중치(Weight): AI 모델이 학습을 통해 저장한 지식. "
        "GPT-4 = 약 1.8조(1.8T) 개의 숫자. 이 숫자들을 매 토큰 생성마다 메모리에서 읽어야 함.")
    _quant(doc,
        "계산 예제 — GPT-4 추론 1토큰 생성 시 메모리 읽기:\n"
        "  파라미터 1.8T × FP16(2바이트) = 3.6TB 데이터를 읽어야 함\n"
        "  DDR5(89 GB/s)로: 3,600GB ÷ 89 GB/s = 40.4초 → 실사용 불가\n"
        "  HBM3e(1,229 GB/s)로: 3,600GB ÷ 1,229 GB/s = 2.9초 → 가능\n"
        "  (실제는 MoE 구조·캐시 최적화로 더 짧아지나, 대역폭이 핵심임을 보여줌)")

    _add_heading(doc, "1.3 메모리 벽(Memory Wall) 개념", 2)
    _concept(doc,
        "Memory Wall(Patterson & Hennessy 1994): CPU/GPU의 연산 속도 발전은 "
        "매년 50%+씩 향상되는 반면, 메모리 대역폭 발전은 연 25% 수준에 그쳐 "
        "격차가 벌어지는 현상. AI 시대에 이 문제가 극단적으로 심화됨.")
    _marketing(doc,
        "마케팅 시사점: 메모리 벽 = 'AI 성능 확장의 영구적 제약'. "
        "이 제약을 가장 잘 해소하는 HBM 공급사(SK Hynix)가 구조적 프리미엄을 누림.")

    _add_refs(doc, ["Patterson_Memory_Wall", "Kaplan_Scaling", "MLCommons_Bench"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 2장. HBM 제조 기술 원리 (마케터를 위한 기술 기초)
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "2장. HBM 제조 기술 원리")

    doc.add_paragraph(
        "마케팅 담당자가 HBM 기술을 알아야 하는 이유: "
        "공급 제약의 근본 원인이 기술에 있기 때문입니다. "
        "기술 난이도가 높을수록 진입장벽이 높고 → 가격 결정력이 강해집니다."
    )

    _add_heading(doc, "2.1 일반 DRAM vs HBM 구조 차이", 2)
    _concept(doc,
        "일반 DRAM: 평면(2D) 구조. 칩 한 장이 메모리 보드 위에 수평으로 배열. "
        "신호선이 길어서 속도에 한계가 있음. (마치 긴 복도 양쪽에 문이 있는 구조)")
    _concept(doc,
        "HBM: 수직(3D) 적층 구조. 여러 장의 DRAM 다이(Die)를 수직으로 쌓고, "
        "TSV(관통 실리콘 비아)라는 수직 전선으로 연결. "
        "(마치 엘리베이터로 연결된 다층 창고 — 이동 거리가 극단적으로 짧아짐)")

    struct_rows = [
        ["적층 구조",   "평면(1층)",         "수직 적층(4~12층, Hi)"],
        ["연결 방식",   "PCB 배선 (수mm)",    "TSV 수직 비아 (수μm, 1000배 짧음)"],
        ["대역폭",      "~89 GB/s (DDR5)",   "~1,229 GB/s (HBM3e)"],
        ["용량/패키지", "최대 수GB",          "8~36GB/스택 (다이수 × 스택 크기)"],
        ["전력 효율",   "기준",              "30% 이상 효율적 (짧은 신호 경로)"],
        ["단가",        "$3-7/GB",           "$14-24/GB (대역폭 프리미엄)"],
        ["주요 제조사", "삼성·SK·Micron 모두", "SK Hynix 52%, Samsung 38%, Micron 10%"],
    ]
    _add_table(doc, ["항목", "일반 DRAM", "HBM"], struct_rows)

    _add_heading(doc, "2.2 TSV(관통 실리콘 비아) 이해하기", 2)
    _concept(doc,
        "TSV(Through-Silicon Via): 실리콘 칩을 수직으로 관통하는 미세 전도성 구멍. "
        "지름 약 5-10μm (사람 머리카락 굵기의 1/10). "
        "이 구멍들을 통해 적층된 DRAM 다이들이 서로 데이터를 주고받습니다.")
    _marketing(doc,
        "TSV는 HBM의 핵심 경쟁 우위이자 진입장벽의 원천입니다. "
        "TSV 드릴링·충전·정렬 기술을 습득하는 데 수년이 걸리며, "
        "SK Hynix는 2013년부터 양산을 시작해 약 10년의 기술 격차를 보유합니다.")
    _marketing(doc,
        "삼성이 HBM3e NVIDIA 인증을 통과 못 하는 주요 이유 중 하나: "
        "TSV 정렬 정확도(Overlay)가 SK Hynix 대비 낮아 수율(Yield)이 떨어지기 때문.")

    _add_heading(doc, "2.3 CoWoS 패키징 — HBM이 GPU에 붙는 방법", 2)
    _concept(doc,
        "HBM을 만들어도 GPU(NVIDIA B200 등)에 물리적으로 연결하려면 "
        "'CoWoS(Chip-on-Wafer-on-Substrate)' 패키징 공정이 필요합니다. "
        "TSMC가 독점 제공하는 이 공정은 GPU 다이와 HBM 스택을 "
        "인터포저(Interposer) 위에 나란히 배치하고 수천 개의 범프(Bump)로 연결합니다.")
    _marketing(doc,
        "CoWoS = 'GPU와 HBM의 결혼식장'. TSMC만 이 공간을 제공할 수 있음. "
        "HBM이 아무리 많아도 CoWoS 캐파(결혼식장 수)가 부족하면 GPU 출하량이 막힙니다. "
        "이것이 HBM과 CoWoS가 '공동 병목(Co-bottleneck)'이 되는 이유입니다.")

    cowos_rows = [
        ["2023", "35,000 wpm", "GPU 50만 개 상당"],
        ["2024", "50,000 wpm", "GPU 80만 개 상당"],
        ["2025", "85,000 wpm", "GPU 180만 개 상당"],
        ["2026E","120,000 wpm","GPU 250만 개 상당 (추정)"],
    ]
    _add_table(doc, ["연도", "TSMC CoWoS 월 캐파(WPM)", "GPU 패키징 가능 물량(추정)"], cowos_rows)
    _risk(doc,
        "CoWoS 리드타임 18개월: 캐파를 늘리려면 18개월이 필요합니다. "
        "즉 오늘 TSMC가 투자를 결정해도 GPU 공급에 반영되는 것은 1.5년 후. "
        "이 기간이 SK Hynix·TSMC 주주의 초과 수익 구간입니다.")

    _add_refs(doc, ["SKHynix_HBM_Blog", "TSMC_CoWoS_2024", "SemiAnalysis_HBM"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3장. HBM 시장 구조 (공급·수요·가격)
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "3장. HBM 시장 구조 — 공급·수요·가격 동학")

    doc.add_paragraph(
        "기술 원리를 이해했으니 이제 시장 수치를 분석합니다. "
        "반도체 마케팅 담당자에게 가장 중요한 질문: "
        "'누가, 얼마나 만들어서, 얼마에 파는가?'"
    )

    _add_heading(doc, "3.1 시장 규모 및 성장 (TAM)", 2)
    tam_rows = [
        ["2019", "$2.5B",  "-",    "HBM2 초기, 주로 슈퍼컴 용도"],
        ["2021", "$4.5B",  "+80%", "HBM2e, A100 GPU 수요"],
        ["2022", "$7.0B",  "+56%", "H100 예약 주문 시작"],
        ["2023", "$14.0B", "+100%","H100 대량 출하, ChatGPT 효과"],
        ["2024", "$22.0B", "+57%", "HBM3e 주력, H200·B200 전환"],
        ["2025E","$35.0B", "+59%", "GB200 NVL72 램프, HBM3e 풀가동"],
        ["2026E","$55.0B", "+57%", "HBM4 진입, B200 Ultra 대응"],
        ["2027E","$80.0B", "+45%", "HBM4e 개발, 에지 AI 수요 추가"],
    ]
    _add_table(doc, ["연도", "시장 규모", "YoY 성장", "핵심 드라이버"], tam_rows)
    _marketing(doc,
        "HBM TAM은 2019-2027년 32배 성장 예상. "
        "이는 반도체 역사상 가장 빠른 세그먼트 성장 중 하나입니다. "
        "비교: DRAM 전체 시장 연 성장률 5-10%.")

    _add_heading(doc, "3.2 공급사별 점유율 (Market Share)", 2)
    share_rows = [
        ["2022", "45%", "45%", "10%", "SK·Samsung 대등; HBM2e 주력"],
        ["2023", "48%", "42%", "10%", "SK Hynix HBM3 출시 우위"],
        ["2024", "52%", "38%", "10%", "SK HBM3e NVIDIA 독점 공급; Samsung 인증 지연"],
        ["2025E","55%", "35%", "10%", "SK HBM3e GB200 전용; Micron 미약한 증가"],
        ["2026E","52%", "38%", "10%", "Samsung HBM3e 재진입 가능; SK HBM4 전환"],
    ]
    _add_table(doc, ["연도", "SK Hynix", "Samsung", "Micron", "비고"], share_rows)
    _marketing(doc,
        "핵심 인사이트: SK Hynix는 기술 리드로 점유율을 늘렸지만, "
        "Samsung이 언제 인증을 통과하느냐가 2026년 이후 시장 구도를 결정합니다. "
        "'Samsung 인증 뉴스'가 투자 카탈리스트가 될 수 있습니다.")

    _add_heading(doc, "3.3 HBM 가격 구조 (ASP & 마진)", 2)
    _concept(doc,
        "HBM의 가격은 일반 DRAM과 구조가 다릅니다. "
        "일반 DRAM: 스팟(현물) 시장 가격이 분기마다 변동. "
        "HBM: 장기계약(LTA, Long-Term Agreement) 기반. "
        "NVIDIA·Google·Amazon이 6-12개월 선행 발주 → 가격 안정성 높음.")
    asp_rows = [
        ["HBM2e",    "$8/GB",     "~70%",      "A100 GPU 탑재; 현재 단종 전환"],
        ["HBM3",     "$14/GB",    "~75%",       "H100 GPU; 점차 HBM3e로 대체"],
        ["HBM3e 8Hi", "$18/GB",   "~78%",       "H200·B200 GPU; 현재 주력"],
        ["HBM3e 12Hi", "$22/GB",  "~80%",       "GB200 NVL72 랙용; 최고 ASP"],
        ["HBM4 (25~)", "~$24+/GB","~80%+",      "Blackwell Ultra; 대역폭 2TB/s"],
    ]
    _add_table(doc, ["제품", "ASP (추정)", "SK Hynix 영업이익률", "비고"], asp_rows)
    _marketing(doc,
        "HBM3e 마진 ~80%는 반도체 역사상 전례 없는 수준입니다. "
        "일반 DRAM 마진이 호황기 30-40%임을 감안하면, "
        "HBM은 완전히 다른 경제학 구조입니다. "
        "이는 공급 제약 + 기술 독점 + LTA 구조가 결합된 결과입니다.")
    _risk(doc,
        "ASP 리스크: Samsung HBM3e 대량 인증 시 ASP 하락 압력. "
        "현재 SK Hynix의 ASP 프리미엄은 '독점적 지위' 덕분. "
        "경쟁 심화 시 마진이 60% 수준으로 수렴할 수 있습니다.")

    _add_refs(doc, ["Omdia_HBM_2024", "TrendForce_HBM", "Bernstein_HBM", "SKHynix_IR_2024"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 4장. HBM 수요 정량 모델
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "4장. HBM 수요 정량 모델 — 직접 계산하기")

    doc.add_paragraph(
        "마케터가 HBM 수요를 직접 산출할 수 있어야 합니다. "
        "이 모델은 GPU 출하량에서 HBM 총 수요를 역산하는 Bottom-up 방식입니다."
    )

    _add_heading(doc, "4.1 GPU당 HBM 탑재량 변화", 2)
    hbm_per_gpu = [
        ["NVIDIA A100 (2020)", "HBM2e × 5스택", "40GB", "~150 GB/s × 5 = 750 GB/s"],
        ["NVIDIA H100 (2022)", "HBM3 × 5스택",  "80GB", "~819 GB/s"],
        ["NVIDIA H200 (2024)", "HBM3e × 6스택", "141GB","~4.8 TB/s"],
        ["NVIDIA B200 (2025)", "HBM3e × 8스택", "192GB","~8.0 TB/s"],
        ["NVIDIA GB200 NVL72 랙", "B200 × 72개", "13.8TB",  "랙 단위 HBM 총량"],
    ]
    _add_table(doc, ["GPU 모델", "HBM 구성", "HBM 용량", "총 대역폭"], hbm_per_gpu)
    _marketing(doc,
        "GB200 NVL72 랙 1개 = HBM 13.8TB. "
        "H100 서버 1개(80GB) 대비 172배 HBM 탑재. "
        "하이퍼스케일러가 GB200로 전환할수록 HBM 수요가 기하급수적으로 증가합니다.")

    _add_heading(doc, "4.2 2026년 HBM 총 수요 산출 (Bottom-up)", 2)
    _quant(doc,
        "전제:\n"
        "  - 2026년 NVIDIA GPU 총 출하 예상: ~1,200만 개 (H100+H200+B200+GB200 합산)\n"
        "  - 믹스: H100/H200 35%, B200 55%, GB200 랙 10%\n\n"
        "계산:\n"
        "  H100/H200: 420만 개 × 80~141GB = 평균 110GB × 420만 = 462PB\n"
        "  B200: 660만 개 × 192GB = 1,267PB\n"
        "  GB200(랙 20만 개): 20만 × 72 × 192GB = 2,765PB\n"
        "  합계: ~4,494PB ≈ 4.4 EB\n\n"
        "공급 상한(SK Hynix + Samsung + Micron, 2026E):\n"
        "  총 HBM WPM: (45K + 38K + 12K) = 95K wpm × 12개월 × 수율 75% × 50GB/wafer\n"
        "  ≈ 95,000 × 12 × 0.75 × 50GB = 42.75M GB = 42.75PB/월 × 12 = 513PB/년\n\n"
        "⚠️ 수요(4,494PB) >> 공급(513PB): 극심한 공급 부족! "
        "(단, 실제 GPU 가동 기반 HBM 수요는 신규 출하 기준이며, "
        "설치베이스 전체를 매년 교체하는 것이 아님 — 순증 수요 기준으로 봐야 함)")

    _add_heading(doc, "4.3 수요-공급 갭 분석 (Simplified)", 2)
    gap_rows = [
        ["2023", "HBM3", "~150PB 수요", "~80PB 공급", "공급 부족 ~70PB"],
        ["2024", "HBM3e", "~300PB 수요", "~200PB 공급", "공급 부족 ~100PB"],
        ["2025E","HBM3e", "~600PB 수요", "~450PB 공급", "공급 부족 ~150PB"],
        ["2026E","HBM3e/4","~900PB 수요","~650PB 공급", "공급 부족 점진 완화"],
        ["2027E","HBM4",   "~1,200PB 수요","~1,000PB 공급","균형 접근"],
    ]
    _add_table(doc, ["연도", "주력 제품", "수요 추정", "공급 추정", "갭"], gap_rows)
    _risk(doc,
        "이 수치는 분석가 추정치이며 정확한 공시 데이터가 아닙니다. "
        "실제 공급-수요 균형은 NVIDIA 출하량 변화, HBM 수율 개선, "
        "비(非)NVIDIA AI 가속기 수요에 따라 달라질 수 있습니다.")

    _add_refs(doc, ["NVIDIA_GTC_B200", "Omdia_HBM_2024", "TrendForce_HBM", "Bernstein_HBM"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 5장. 경쟁사 심화 분석
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "5장. 경쟁사 심화 분석 — SK Hynix vs Samsung vs Micron")

    doc.add_paragraph(
        "HBM 시장은 3사 과점이지만, 실질적으로는 SK Hynix의 '기술 독점' 구조입니다. "
        "각 플레이어의 현황, 강점, 리스크를 마케팅 관점으로 분석합니다."
    )

    _add_heading(doc, "5.1 SK Hynix — 현재 챔피언", 2)
    doc.add_paragraph(
        "SK Hynix가 HBM3e 시장을 주도하는 구체적 이유:\n\n"
        "① 10년 선행 투자: 2013년 최초 HBM 양산 시작. 삼성·Micron 대비 3-5년 앞섬.\n"
        "② NVIDIA 전용 공급 계약: HBM3e 생산의 80%+가 NVIDIA 할당. "
        "Google/AWS 등은 나머지로 분배.\n"
        "③ 수율(Yield) 우위: HBM3e 8-Hi 수율 약 85% (삼성 추정 65-70% 대비). "
        "수율 10% 차이 = 같은 웨이퍼에서 생산량 ~15% 차이.\n"
        "④ HBM4 선점: 2025년 하반기 양산 시작. 대역폭 2,000GB/s, "
        "GB200 Ultra(Blackwell Ultra) GPU 탑재 예정."
    )
    _marketing(doc,
        "SK Hynix 영업전략 핵심: NVIDIA와의 JDP(Joint Development Program). "
        "GPU 설계 초기 단계부터 HBM 스펙을 함께 정의 → "
        "경쟁사가 스펙을 공개 후 따라갈 때 SK Hynix는 이미 양산 중. "
        "이것이 'Design Win'의 반도체 메모리 버전입니다.")

    _add_heading(doc, "5.2 Samsung — 리스크 & 기회", 2)
    doc.add_paragraph(
        "Samsung이 HBM3e NVIDIA 인증에 고전하는 구체적 원인:\n\n"
        "① 발열 문제: Samsung HBM3e의 작동 온도가 SK Hynix 대비 높아 "
        "GB200 NVL72 랙의 열 설계(Thermal Design) 기준을 충족 못함.\n"
        "② 수율 격차: TSV 정렬 정밀도가 낮아 8-Hi 스태킹 시 불량률 높음.\n"
        "③ TSMC CoWoS 물량 부족: Samsung이 인증 통과해도 CoWoS 패키징 슬롯이 "
        "SK Hynix 할당으로 가득 차 있음.\n\n"
        "기회 요인:\n"
        "① HBM4 에서 재도전: 새로운 세대는 기술 격차가 초기화될 수 있음.\n"
        "② 자체 파운드리(SF3/SF2) CoWoS 대안 개발 중 (TSMC 의존도 탈피 시도).\n"
        "③ 비NVIDIA 고객(Google TPU, Amazon Trainium) 확보 가능."
    )
    _marketing(doc,
        "Samsung 투자 관점: '카탈리스트 주도형' 포지션. "
        "'Samsung HBM3e NVIDIA 인증 완료' 뉴스가 나오면 "
        "주가가 급등할 수 있는 비대칭 기회. "
        "단, 인증 지연이 장기화될 경우 HBM4 경쟁에서도 뒤처질 위험이 있음.")

    _add_heading(doc, "5.3 Micron — 지정학적 카드", 2)
    doc.add_paragraph(
        "Micron의 HBM 포지셔닝 전략:\n\n"
        "① 점유율 목표: 현재 10% → 2026년 20%+ 목표. '다크호스' 전략.\n"
        "② 지정학적 강점: 미국 CHIPS Act 수혜. "
        "NVIDIA·US 하이퍼스케일러들이 '미국산 HBM' 비율을 높이려는 규제 압력 대응.\n"
        "③ Idaho(보이시) HBM 전용 팹 건설 중: $15B 투자, 2026-2027 양산 목표.\n"
        "④ 수율 개선 속도: Micron HBM3e 8-Hi 수율이 분기마다 5-10% 개선 중."
    )
    _marketing(doc,
        "Micron의 세일즈 포인트: '미국에서 만든 HBM'. "
        "특히 US 정부 조달, 방산 관련 AI 인프라에서 유일한 선택지. "
        "이 '지정학적 모트'는 기술보다 장기적으로 지속될 수 있습니다.")

    comp_summary = [
        ["SK Hynix",  "NVIDIA 전용, 수율 우위, 10년 선행",   "HBM4 전환 리스크, Korea 지정학",   "Strong Buy"],
        ["Samsung",   "규모·재무 여력, HBM4 재도전",         "인증 지연, 발열 문제, CoWoS 부족", "Watch (카탈리스트 대기)"],
        ["Micron",    "미국산 지정학 강점, CHIPS Act 보조금", "기술 격차, 수율 램프 리스크",       "Buy (중장기)"],
    ]
    _add_table(doc, ["기업", "핵심 강점", "핵심 리스크", "투자 시그널"], comp_summary)

    _add_refs(doc, ["SKHynix_HBM_Blog", "SKHynix_IR_2024", "Micron_HBM3E",
                    "Omdia_HBM_2024", "Bernstein_HBM"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 6장. HBM 투자 프레임워크 (마케터 → 투자자 관점 전환)
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "6장. HBM 투자 프레임워크")

    doc.add_paragraph(
        "반도체 마케팅 전문가가 HBM 투자 분석에서 가장 유리한 점: "
        "제품 사이클·설계 승인(Design Win)·고객사 구매 패턴을 직관적으로 이해한다는 것입니다. "
        "이 섹션은 그 지식을 투자 프레임워크로 전환합니다."
    )

    _add_heading(doc, "6.1 HBM 투자 수익의 원천 (3요소)", 2)
    return_rows = [
        ["① 물량 증가\n(Volume Growth)",
         "AI GPU 출하 증가 → HBM 탑재량 증가",
         "H100→B200 전환 시 GPU당 HBM 2.4배 증가\n(80GB→192GB)"],
        ["② ASP 상승\n(Price Increase)",
         "세대 전환 시 HBM당 ASP 20-30% 상승",
         "HBM3→HBM3e: $14→$18/GB (+29%)\nHBM3e→HBM4: $18→$24/GB (+33%)"],
        ["③ 점유율 유지\n(Share of Wallet)",
         "SK Hynix의 NVIDIA 공급 비중 유지/확대",
         "현재 80%+ SK Hynix → HBM4 에서도 유지 목표"],
    ]
    _add_table(doc, ["수익 원천", "메커니즘", "구체적 수치"], return_rows)

    _quant(doc,
        "SK Hynix HBM 매출 성장 추정 (간단 계산):\n"
        "  2024: $22B 시장 × 52% 점유율 = ~$11.4B\n"
        "  2026E: $55B 시장 × 52% 점유율 = ~$28.6B\n"
        "  2024→2026 성장: +$17.2B (+151%) — 2년 만에 2.5배\n"
        "  전사 영업이익 영향: HBM 영업이익률 ~80% → 약 +$13.8B 이익 증가 기여")

    _add_heading(doc, "6.2 병목 투자 매트릭스 (가동률 × 리드타임)", 2)
    doc.add_paragraph(
        "병목 레이어 투자의 핵심 원칙: '가동률이 높고 해소 리드타임이 길수록 투자 매력이 높다.'"
    )
    matrix_rows = [
        ["HBM",         "92%", "높음 ★★★", "18개월", "★★★ (최상)", "SK Hynix, Micron"],
        ["CoWoS",       "88%", "높음 ★★★", "12-18개월","★★★",      "TSMC (패키징 부문)"],
        ["Power_DC",    "85%", "높음 ★★★", "24-36개월","★★★",      "Vertiv, GE Vernova"],
        ["Networking",  "78%", "중간 ★★",  "6-12개월", "★★",        "Broadcom, Marvell"],
        ["GPU",         "87%", "중간 ★★",  "9-12개월", "★★",        "NVIDIA (선반영)"],
        ["DRAM",        "55%", "낮음 ★",   "6개월",    "★",         "Cycle 하락 주의"],
        ["SSD",         "40%", "낮음 ★",   "3개월",    "★",         "공급 과잉 주의"],
    ]
    _add_table(doc,
        ["레이어", "현 가동률", "병목 심각도", "해소 리드타임", "투자 매력도", "수혜 기업"],
        matrix_rows)

    _add_heading(doc, "6.3 카탈리스트 타임라인 (2026-2027)", 2)
    catalyst_rows = [
        ["2026 Q1-Q2", "HBM4 양산 확대",
         "SK Hynix HBM4 물량 증가 확인 시 ASP 프리미엄 지속 확인"],
        ["2026 Q2-Q3", "Samsung HBM3e NVIDIA 인증 여부",
         "통과 시: 경쟁 심화 → SK Hynix ASP 하락 압력\n실패 시: SK Hynix 독점 연장"],
        ["2026 Q3-Q4", "GB200 NVL72 대규모 배포",
         "랙당 HBM 13.8TB × 대량 출하 = HBM 수요 폭발 확인"],
        ["2026 Q4",    "TSMC CoWoS 120K wpm 도달",
         "패키징 병목 완화 → HBM 공급 흡수 가속"],
        ["2027 Q1-Q2", "HBM 수요-공급 균형 접근",
         "균형 전환 시 ASP 하락 선행 투자 판단 필요 (매도 시점 모니터링)"],
    ]
    _add_table(doc, ["시기", "카탈리스트", "투자 시사점"], catalyst_rows)
    _marketing(doc,
        "마케팅 담당자 어드밴티지: 고객사(하이퍼스케일러)의 구매 동향, "
        "경쟁사 제품 출시 일정, 산업 컨퍼런스 발표 내용을 조기에 파악 가능. "
        "이 정보를 투자 카탈리스트 타임라인에 연결하면 강력한 edge가 됩니다.")

    _add_refs(doc, ["Goldman_AI_Infra", "Omdia_HBM_2024", "Bernstein_HBM",
                    "TrendForce_HBM", "SKHynix_IR_2024"])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 7장. 핵심 리스크 분석
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "7장. HBM 투자의 핵심 리스크")

    risk_rows = [
        ["Samsung HBM3e 인증 통과",
         "SK Hynix 독점 붕괴 → ASP 하락 20-30%",
         "높음",
         "모니터링: NVIDIA 공급 다변화 뉴스, Samsung IR 발언"],
        ["DeepSeek식 효율화 심화",
         "모델 효율↑ → GPU당 HBM 수요 감소\n(단, Jevons Paradox로 총수요 증가 가능)",
         "중간",
         "단기 주가 변동 리스크; 장기는 오히려 수요 증폭"],
        ["HBM4 수율 문제",
         "신규 세대 전환 시 초기 수율 저하 → 공급 부족 심화",
         "중간",
         "SK Hynix Q3/Q4 컨퍼런스 콜 수율 발언 주시"],
        ["지정학 리스크",
         "미-중 갈등 → 한국 반도체에 대한 미국/중국 이중 압력",
         "낮음~중간",
         "수출 규제 확대 여부, 한국 외교 동향 모니터링"],
        ["수요 사이클 전환",
         "2027년 이후 HBM 공급 과잉 전환 시 ASP·마진 급락",
         "중간~높음 (장기)",
         "CoWoS 캐파 증가 속도 vs 수요 성장 속도 격차 추적"],
        ["NVLINK 외 대안 아키텍처",
         "광(Photonic) 인터커넥트 또는 CXL 메모리가 HBM 대체 가능성",
         "낮음 (2028년 이후)",
         "현재 기술 성숙도 낮아 단기 위협 미미"],
    ]
    _add_table(doc, ["리스크", "영향", "발생 가능성", "모니터링 포인트"], risk_rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 8장. 복습 문제
    # ══════════════════════════════════════════════════════════
    _add_heading(doc, "8장. 복습 문제 — 능동적 학습")

    doc.add_paragraph(
        "페인만 기법: 아래 질문에 '상대방에게 설명하듯' 답할 수 있으면 이해한 것입니다. "
        "막히는 부분이 있으면 해당 장으로 돌아가서 다시 읽으세요."
    )

    questions = [
        # (질문, 힌트, 유형)
        (
            "HBM이 일반 DDR5보다 AI 추론에 유리한 이유를 "
            "'대역폭'과 '메모리 벽' 개념을 사용해 설명하세요.",
            "1장 참고. GPT-4 1토큰 생성 시 필요한 데이터 읽기 계산 예제를 활용하세요.",
            "개념 이해 (기술 원리)"
        ),
        (
            "TSV(관통 실리콘 비아)가 없으면 HBM이 왜 불가능한지 설명하고, "
            "삼성이 HBM3e 인증에 고전하는 기술적 이유와의 연결점을 설명하세요.",
            "2장 참고. '신호 경로 길이'와 '수율(Yield)'의 관계를 생각해보세요.",
            "기술 심화 + 경쟁 분석"
        ),
        (
            "2026년 NVIDIA B200 GPU 출하량이 600만 개로 가정될 때, "
            "B200 GPU가 생성하는 HBM3e 수요를 GB 단위로 계산하세요.",
            "4장 계산 예제 참고. B200 1개당 HBM3e 192GB.",
            "정량 계산 (비즈니스 수학)"
        ),
        (
            "SK Hynix의 HBM 영업이익률이 ~80%인 이유를 "
            "시장 구조·기술 독점·계약 방식 3가지 관점에서 설명하세요.",
            "3장과 6장 참고. LTA(장기계약), 수율 우위, NVIDIA JDP 관계를 활용하세요.",
            "비즈니스 분석 (마케터 강점 영역)"
        ),
        (
            "Samsung HBM3e NVIDIA 인증 통과 뉴스가 나왔을 때 "
            "SK Hynix 주가에 미칠 영향을 예측하고 근거를 제시하세요. "
            "(Short-term vs Long-term 구분)",
            "5장 참고. ASP 하락 압력 vs. 총 시장 확대 효과를 모두 고려하세요.",
            "투자 분석 (시나리오 사고)"
        ),
        (
            "CoWoS 캐파가 HBM 수요에 독립적인 제약이 되는 이유를 설명하고, "
            "이것이 TSMC에게 어떤 가격 결정력을 부여하는지 분석하세요.",
            "2장 CoWoS 섹션 참고. '결혼식장' 비유를 사용해 설명해보세요.",
            "공급망 분석 (병목 연쇄)"
        ),
        (
            "당신이 SK Hynix의 HBM 제품 마케팅 담당자라면, "
            "NVIDIA 구매 담당자에게 HBM4의 가격 프리미엄($24+/GB)을 정당화하기 위해 "
            "어떤 수치와 논리를 사용하겠습니까?",
            "4장의 GPU당 HBM 용량 증가, 6장의 ROI 계산을 활용하세요. "
            "'대역폭 1GB/s당 가격'으로 비교해보는 것도 방법입니다.",
            "실무 적용 (마케터 관점)"
        ),
    ]
    _quiz_section(doc, questions)

    # ── 정답 가이드 ──────────────────────────────────────────
    _add_heading(doc, "정답 가이드 (핵심 포인트)", 2)
    answers = [
        ("Q1 핵심",
         "DDR5=89GB/s vs HBM3e=1,229GB/s (14배). "
         "GPT-4(1.8T 파라미터)는 토큰 1개 생성에 3.6TB 데이터 읽기 필요. "
         "DDR5로는 40초, HBM3e로는 2.9초 → AI 실시용 불가 vs 가능."),
        ("Q2 핵심",
         "TSV가 없으면 다이 간 신호선이 수mm → 속도 한계. "
         "Samsung의 HBM3e 수율 저하: TSV 정렬 정확도(Overlay) 부족 → "
         "8-Hi 스태킹 시 불량 증가 → 수율 65-70% (SK Hynix 85% 대비)."),
        ("Q3 핵심",
         "B200 600만 개 × 192GB = 1,152,000,000 GB = 1,152 PB = 1.15 EB. "
         "(비교: HBM 총 연간 공급 ~600-700 PB 수준 → 심각한 부족)"),
        ("Q4 핵심",
         "① 공급 제약(희소성) → 가격 협상력. "
         "② TSV 수율 우위 → 단위 원가 낮음 → 마진 높음. "
         "③ LTA(장기계약) → 가격 변동성 없음, 안정적 고마진 유지."),
        ("Q5 핵심",
         "단기: 주가 하락 (독점 프리미엄 소멸 우려). "
         "장기: 중립~약간 부정적 (SK Hynix 점유율 일부 잠식, "
         "단 총 시장 확대로 절대 매출은 유지 가능). "
         "투자 포인트: 인증 통과 후 SK Hynix ASP 실제 하락 폭 확인이 중요."),
        ("Q6 핵심",
         "CoWoS = HBM+GPU를 연결하는 유일한 공정 → TSMC 독점. "
         "HBM 공급이 늘어도 CoWoS 캐파가 없으면 GPU 출하 불가. "
         "TSMC는 이 독점으로 CoWoS ASP를 20%+ 인상 (2024년 실제 인상)."),
        ("Q7 핵심",
         "논리: ① 대역폭/$ 효율 → HBM4는 HBM3e 대비 대역폭 63% 증가, "
         "가격 33% 증가 → '단위 대역폭 비용' 실질 감소. "
         "② 서버 ROI: HBM4 탑재 시 추론 속도 +50% → "
         "같은 서버 투자로 더 많은 토큰 처리 → NVIDIA GPU당 ROI 개선."),
    ]
    for key, ans in answers:
        p = doc.add_paragraph()
        r = p.add_run(f"[{key}] ")
        r.bold = True
        r.font.size = Pt(9)
        r2 = p.add_run(ans)
        r2.font.size = Pt(9)

    doc.add_page_break()

    # ── 종합 참고문헌 ──────────────────────────────────────────
    _add_heading(doc, "종합 참고문헌")
    ref_rows = [(n, url, desc, yr) for n, url, desc, yr in REFS.values()]
    _add_table(doc, ["출처명", "URL", "설명", "연도"], ref_rows)

    # ── 저장 ─────────────────────────────────────────────────
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today_str = date.today().strftime("%Y%m%d")
    out = OUTPUT_DIR / f"study_hbm_deep_dive_{today_str}.docx"
    doc.save(str(out))
    print(f"  [HBM Study] 생성 완료: {out}")
    print(f"  [HBM Study] 8장 구성, {len(REFS)}개 참고자료, 7문항 복습 문제")
    return str(out)


def run():
    return {"word_path": build()}


if __name__ == "__main__":
    result = run()
    print(f"\n출력 파일: {result['word_path']}")
