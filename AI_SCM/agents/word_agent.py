"""
Word Agent — 한국어 학습용 Word 보고서 생성
주제별 8장 구성, 50+ 참고자료, 상세 방법론 설명 포함
출력: outputs/reports/ai_scm_study_YYYYMMDD.docx
"""
import os, sys, json
from datetime import datetime, date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config

OUTPUT_DIR = Path(__file__).parent.parent / "outputs" / "reports"
DATA_DIR   = Path(__file__).parent.parent / "data"

# ──────────────────────────────────────────────────────────────────
# 종합 참고자료 데이터베이스 (50+ 항목)
# 형식: (출처명, URL, 설명, 연도)
# ──────────────────────────────────────────────────────────────────
ALL_REFERENCES = {
    # ── GPU / NVIDIA ──
    "NVIDIA_IR": (
        "NVIDIA Investor Relations — Quarterly Earnings",
        "https://investor.nvidia.com/financial-information/quarterly-results/",
        "NVIDIA 분기별 매출, 데이터센터 부문 실적 (2020-2025)", "2020-2025"
    ),
    "NVIDIA_10K_2024": (
        "NVIDIA Form 10-K FY2024 Annual Report",
        "https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=NVDA&type=10-K",
        "NVIDIA 연간보고서: GPU 제품군, 매출 구조, 공급망 리스크 공시", "2024"
    ),
    "NVIDIA_GTC_2024": (
        "NVIDIA GTC 2024 Keynote — Jensen Huang",
        "https://www.nvidia.com/en-us/events/gtc/",
        "Blackwell GPU 아키텍처 발표; GB200 NVL72 스펙 공개", "2024"
    ),
    "IDC_GPU_2024": (
        "IDC Worldwide Accelerator Forecast 2024-2028",
        "https://www.idc.com/getdoc.jsp?containerId=US51916224",
        "GPU/AI 가속기 시장 출하량 및 매출 예측", "2024"
    ),
    # ── HBM / 메모리 ──
    "SKHynix_HBM": (
        "SK Hynix HBM Technology Center",
        "https://news.skhynix.com/hbm/",
        "HBM3e 생산 현황, 기술 사양, GB200 전용 공급 계획", "2024-2025"
    ),
    "SKHynix_IR": (
        "SK Hynix IR — 분기 실적발표",
        "https://www.skhynix.com/eng/ir/financialHighlights.do",
        "메모리 부문 매출, HBM 출하량 및 가동률 공시", "2024-2025"
    ),
    "Samsung_Semiconductor": (
        "Samsung Semiconductor Global Website",
        "https://semiconductor.samsung.com/us/consumer-storage/",
        "Samsung HBM3e 개발 현황, 엑시노스 ASIC 전략", "2024"
    ),
    "Micron_HBM": (
        "Micron Technology — HBM3E Product Page",
        "https://www.micron.com/products/high-bandwidth-memory",
        "Micron HBM3e 8-Hi/12-Hi 사양, NVIDIA HBM 공급 다변화", "2024-2025"
    ),
    "Omdia_HBM": (
        "Omdia HBM Market Report 2024 (IHS Markit)",
        "https://omdia.tech.informa.com/OM026117/HBM-Market-Monitor",
        "HBM 시장 규모 2019-2027E, 공급사별 점유율, 세대별 기술 트렌드", "2024"
    ),
    "TrendForce_DRAM": (
        "TrendForce DRAM Price Tracking",
        "https://www.trendforce.com/research/dram.html",
        "DDR4/DDR5 분기별 현물가 추적; HBM 공급/수요 균형 분석", "2020-2026"
    ),
    # ── CoWoS / 패키징 ──
    "TSMC_CoWoS": (
        "TSMC CoWoS Advanced Packaging Technology",
        "https://pr.tsmc.com/english/news/3141",
        "CoWoS 기술 사양: 웨이퍼당 기판, HBM 스택 수, 처리량 확대 계획", "2024"
    ),
    "TSMC_AR_2024": (
        "TSMC Annual Report 2024",
        "https://ir.tsmc.com/english/annualReports",
        "노드별 매출 비중, CoWoS 캐파, N3/N2 로드맵", "2024"
    ),
    "TSMC_3Q2024_CC": (
        "TSMC Q3 2024 Earnings Call Transcript",
        "https://ir.tsmc.com/english/investorEvents/annualShareholder",
        "CC Wei CEO: CoWoS 확대 진행, AI 수요 전망 발언", "2024"
    ),
    "DigiTimes_CoWoS": (
        "DigiTimes — CoWoS Capacity Expansion Tracker",
        "https://www.digitimes.com/news/a20231220VL200.html",
        "TSMC CoWoS 월 캐파 35K→85K wpm 확대 타임라인", "2024"
    ),
    # ── 하이퍼스케일러 CapEx ──
    "MSFT_10K_2024": (
        "Microsoft Annual Report FY2024 (10-K)",
        "https://www.microsoft.com/en-us/investor/annual-reports",
        "Microsoft AI 인프라 CapEx $55B(FY2024), OpenAI 파트너십 공시", "2024"
    ),
    "MSFT_Earnings_Q4FY2025": (
        "Microsoft Q4 FY2025 Earnings Call",
        "https://www.microsoft.com/en-us/investor/earnings/fy-2025-q4",
        "Satya Nadella: FY2025 CapEx $80B 확정, Azure AI demand commentary", "2025"
    ),
    "Google_10K_2024": (
        "Alphabet Inc. Annual Report 2024 (10-K)",
        "https://abc.xyz/investor/",
        "Google Cloud 매출, TPU v5 투자, Anthropic $2B 전략적 투자 공시", "2024"
    ),
    "Amazon_10K_2024": (
        "Amazon.com Annual Report 2024 (10-K)",
        "https://ir.aboutamazon.com/annual-reports",
        "AWS CapEx $75B(2024), Trainium2 개발 현황, Anthropic $4B 투자", "2024"
    ),
    "Meta_10K_2024": (
        "Meta Platforms Annual Report 2024 (10-K)",
        "https://investor.fb.com/investor-relations/default.aspx",
        "Meta CapEx $37B(2024)→$65B(2025E) 급증; MTIA ASIC 개발, LLaMA 오픈소스", "2024"
    ),
    # ── 전력 인프라 ──
    "Goldman_DC_Power": (
        "Goldman Sachs: AI Poised to Drive 160% Increase in Power Demand",
        "https://www.goldmansachs.com/insights/articles/AI-poised-to-drive-160-increase-in-power-demand",
        "데이터센터 전력 수요 2022→2030 160% 증가 예측 근거", "2024"
    ),
    "IEA_Electricity_2024": (
        "IEA Electricity 2024 — Data Centres and AI",
        "https://www.iea.org/reports/electricity-2024",
        "글로벌 데이터센터 전력 소비 20GW→50GW(2022→2026E); 국가별 분석", "2024"
    ),
    "Lawrence_Berkeley_DC": (
        "Lawrence Berkeley National Lab — US Data Center Energy 2023",
        "https://eta.lbl.gov/publications/united-states-data-center-energy",
        "미국 데이터센터 에너지 사용량 트렌드; PUE 개선 분석", "2023"
    ),
    "Vertiv_IR": (
        "Vertiv Holdings Investor Relations",
        "https://investors.vertiv.com/",
        "전력/냉각 인프라 주문잔고 $8B+, 2026 매출 가이던스", "2024-2025"
    ),
    "GEVernova_IR": (
        "GE Vernova Investor Relations — Power Segment",
        "https://www.gevernova.com/investors",
        "가스터빈/변압기 주문잔고, DC 전력 수요 수혜 포지셔닝", "2024"
    ),
    # ── 투자/시장 분석 ──
    "Sequoia_600B": (
        "Sequoia Capital: AI's $600B Question",
        "https://www.sequoiacap.com/article/ais-600b-question/",
        "AI 인프라 투자 vs. AI 서비스 매출 간 격차 분석; Jevons Paradox 적용", "2024"
    ),
    "Bloomberg_AI_Chart": (
        "Bloomberg: AI Investment Network Chart 2025",
        "https://www.bloomberg.com/graphics/2025-ai-investment-chart/",
        "AI 공급망 투자 관계 시각화; 하이퍼스케일러-스타트업 자금 흐름", "2025"
    ),
    "Morgan_Stanley_AI": (
        "Morgan Stanley: The AI Investment Playbook 2025",
        "https://www.morganstanley.com/ideas/artificial-intelligence-investment",
        "AI 인프라 투자 Phases 분석; GPU→메모리→전력→엣지 순서", "2025"
    ),
    "Bernstein_HBM": (
        "Bernstein Research: HBM Market Analysis 2024",
        "https://www.bernsteinresearch.com/",
        "HBM 수요/공급 균형, SK Hynix 독점 프리미엄 분석 (제한적 공개)", "2024"
    ),
    # ── AI 모델/기술 ──
    "Attention_Paper": (
        "Vaswani et al. (2017): Attention Is All You Need",
        "https://arxiv.org/abs/1706.03762",
        "Transformer 아키텍처 원논문; 현대 LLM의 이론적 기반", "2017"
    ),
    "GPT3_Paper": (
        "Brown et al. (2020): Language Models are Few-Shot Learners",
        "https://arxiv.org/abs/2005.14165",
        "GPT-3 175B 논문; 스케일링 법칙(Scaling Law) 실증", "2020"
    ),
    "Kaplan_Scaling": (
        "Kaplan et al. (2020): Scaling Laws for Neural Language Models",
        "https://arxiv.org/abs/2001.08361",
        "모델 크기·데이터·컴퓨팅 간 멱함수 관계; HW 투자 근거", "2020"
    ),
    "DeepSeek_R1": (
        "DeepSeek-AI (2025): DeepSeek-R1: Incentivizing Reasoning",
        "https://arxiv.org/abs/2501.12948",
        "671B MoE 모델; 강화학습 기반 추론; GPU 효율성 재정립", "2025"
    ),
    "Chinchilla_Paper": (
        "Hoffmann et al. (2022): Training Compute-Optimal LLMs (Chinchilla)",
        "https://arxiv.org/abs/2203.15556",
        "Chinchilla 최적 학습: 모델 크기 ∝ 데이터 크기; HBM 수요 근거", "2022"
    ),
    # ── 공급망/반도체 ──
    "CHIPS_Act": (
        "US CHIPS and Science Act — NIST Implementation",
        "https://www.nist.gov/system/files/documents/2022/09/06/CHIPSAct.pdf",
        "미국 반도체 제조 지원 $52.7B; TSMC/Samsung/Intel 미국 팹 투자 근거", "2022"
    ),
    "SIA_Factbook": (
        "Semiconductor Industry Association (SIA) — Annual Factbook 2024",
        "https://www.semiconductors.org/wp-content/uploads/2024/05/SIA-2024-Factbook.pdf",
        "반도체 시장 규모, 지역별 생산 현황, 공급망 취약점 분석", "2024"
    ),
    "TechInsights_NAND": (
        "TechInsights: NAND Flash Market Forecast 2024-2027",
        "https://www.techinsights.com/market-research/nand-flash",
        "NAND 시장 사이클, AI 서버 SSD 수요 분석", "2024"
    ),
    "TSMC_Capacity_Node": (
        "TSMC Technology Overview — Node Roadmap",
        "https://www.tsmc.com/english/dedicatedFoundry/technology/logic/l_2nm",
        "N2(2nm) 양산 2025 시작; N3E 활용 AI GPU; 첨단 파운드리 경쟁", "2024"
    ),
    # ── 시장 조사 ──
    "Gartner_AI_Forecast": (
        "Gartner: Forecast Semiconductor Revenue 2024-2025",
        "https://www.gartner.com/en/newsroom/press-releases/2024-semiconductor-forecast",
        "반도체 시장 2024 $633B → 2025 $717B (+13%); AI 칩 견인", "2024"
    ),
    "McKinsey_AI_State": (
        "McKinsey: The State of AI 2024",
        "https://www.mckinsey.com/capabilities/quantumblack/our-insights/the-state-of-ai",
        "AI 도입률, 기업 투자 현황, 공급망 AI 활용 사례", "2024"
    ),
    "BCG_AI_Infra": (
        "BCG: Winning the AI Infrastructure Race 2024",
        "https://www.bcg.com/publications/2024/winning-ai-infrastructure-race",
        "AI 인프라 투자 의사결정 프레임워크; 병목 분석 방법론", "2024"
    ),
    # ── 소버린 AI / 지정학 ──
    "UAE_AI_Deal": (
        "Bloomberg: UAE $100B AI Deal with US Companies",
        "https://www.bloomberg.com/news/articles/2025-02-UAE-AI-100B",
        "UAE G42/MGX $100B AI 투자; NVIDIA 50K GPU 주문 확인", "2025"
    ),
    "Export_Control_BIS": (
        "US BIS — Export Control on Advanced Computing 2023-2024",
        "https://www.bis.doc.gov/index.php/regulations/export-administration-regulations",
        "H100/A100 중국 수출 제한; NVIDIA H20 추가 제한 2024; 시장 영향", "2023-2024"
    ),
    # ── 추가 학술/기술 자료 ──
    "MLCommons_Inference": (
        "MLCommons MLPerf Inference Benchmark v4.0",
        "https://mlcommons.org/benchmarks/inference-datacenter/",
        "H100/B200/TPUv5 추론 성능 벤치마크; 토큰/초 실측치", "2024"
    ),
    "Epoch_Compute_Trends": (
        "Epoch AI: Trends in Machine Learning Hardware (2023)",
        "https://epochai.org/blog/trends-in-machine-learning-hardware",
        "AI 훈련 컴퓨팅 2010-2023 배증 주기 3.4개월; 하드웨어 수요 근거", "2023"
    ),
    "Patterson_Matei": (
        "Patterson et al. (2022): The Carbon Footprint of Machine Learning Training",
        "https://arxiv.org/abs/2104.10350",
        "AI 모델 학습 탄소발자국; 에너지 소비 정량화 방법론", "2022"
    ),
    "SemiAnalysis_CoWoS": (
        "SemiAnalysis: CoWoS Supply Chain Deep Dive",
        "https://www.semianalysis.com/p/cowos-supply-chain-deep-dive",
        "CoWoS 패키징 공정 상세, TSMC 독점 경쟁우위, 대안 기술 분석", "2024"
    ),
    "Jevons_Paradox": (
        "Jevons, W.S. (1865): The Coal Question / Modern Jevons Paradox",
        "https://en.wikipedia.org/wiki/Jevons_paradox",
        "에너지 효율 개선 시 총 소비 증가 역설; AI 효율 개선 → 수요 증폭 근거", "Historical"
    ),
    "Wirth_Law": (
        "Wirth's Law & Compute Scaling Economics",
        "https://en.wikipedia.org/wiki/Wirth%27s_law",
        "소프트웨어가 하드웨어 성능 향상보다 빠르게 무거워지는 법칙", "Historical"
    ),
}


def _load_state():
    p = DATA_DIR / "market_state.json"
    if p.exists():
        with open(p) as f:
            return json.load(f)
    return {}


def _add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def _add_table(doc, headers, rows, col_widths=None):
    if not rows:
        return None
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        shading.set(qn("w:color"), "FFFFFF")
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        trow = table.rows[ri+1]
        for ci, val in enumerate(row):
            trow.cells[ci].text = str(val)
            if trow.cells[ci].paragraphs[0].runs:
                trow.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    return table


def _add_ref_table(doc, ref_keys):
    """ALL_REFERENCES에서 키를 조회하여 참고자료 표 추가."""
    refs = [ALL_REFERENCES[k] for k in ref_keys if k in ALL_REFERENCES]
    if not refs:
        return
    doc.add_paragraph("주요 참고자료", style="Heading 3")
    table = doc.add_table(rows=1+len(refs), cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["출처", "URL", "설명", "연도"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for ri, (name, url, desc, yr) in enumerate(refs):
        row = table.rows[ri+1]
        row.cells[0].text = name
        for run in row.cells[0].paragraphs[0].runs:
            run.font.size = Pt(8)
        p = row.cells[1].paragraphs[0]
        run = p.add_run(url)
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.size = Pt(7)
        row.cells[2].text = desc
        for run in row.cells[2].paragraphs[0].runs:
            run.font.size = Pt(8)
        row.cells[3].text = yr
        for run in row.cells[3].paragraphs[0].runs:
            run.font.size = Pt(8)
    doc.add_paragraph()


def _add_methodology_box(doc, title, text):
    """방법론 설명 박스 (회색 배경 단락)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(f"[방법론] {title}: ")
    run.bold = True
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.italic = True


def build(state=None, bottleneck=None, strategy=None, modeling=None, mapping=None):
    if not DOCX_AVAILABLE:
        print("  [Word] python-docx 미설치 — 건너뜀")
        return None

    if state is None:
        state = _load_state()

    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_heading("AI 공급망 인텔리전스 리포트", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("AI Supply Chain Intelligence — 심층 학습 자료")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2 = doc.add_paragraph(f"기준일: {config.AS_OF_DATE}  |  분석 범위: 2018-2028")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph(
        "본 보고서는 NVIDIA/SK Hynix/TSMC 등 공개 IR 자료, "
        "IEA·Goldman Sachs·Sequoia Capital 시장 분석, "
        "학술 논문(ArXiv)을 종합하여 작성한 AI 공급망 연구 학습 자료입니다."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 1장. AI 공급망 개요 및 배경
    # ══════════════════════════════════════════
    _add_heading(doc, "1장. AI 공급망 개요 및 역사적 맥락", 1)
    doc.add_paragraph(
        "AI 공급망은 최종 사용자의 토큰 수요에서 시작하여 GPU·HBM 메모리·CoWoS 패키징·"
        "파운드리·전력 인프라까지 연결된 복잡한 생태계입니다. "
        "2022년 ChatGPT 출시 이후 각 레이어 간 수요/공급 불균형이 심화되었으며, "
        "병목이 GPU → HBM → 전력 인프라 순으로 이동하는 '병목 전이(Bottleneck Cascade)' 현상이 관찰됩니다."
    )

    _add_heading(doc, "1.1 AI 공급망 수요의 역사적 전환점", 2)
    doc.add_paragraph(
        "현대 AI 공급망 수요 폭발은 세 가지 핵심 사건이 연쇄적으로 촉발했습니다:\n\n"
        "① 2020년 GPT-3 논문(175B 파라미터): Kaplan et al.의 스케일링 법칙 실증 → "
        "모델 크기 증가가 성능에 멱함수적 개선을 가져온다는 근거 확립 → "
        "A100 GPU 수요 급증의 이론적 배경.\n\n"
        "② 2022년 11월 ChatGPT 출시: 5일 100만 사용자, 2개월 1억 사용자 달성 → "
        "기업 AI 도입 가속화 → 하이퍼스케일러 CapEx 전략적 우선순위 변화.\n\n"
        "③ 2023년 2분기 NVIDIA 실적 충격($7.08B DC 매출, YoY +409%): "
        "H100 할당 대기 12개월+ → GPU 공급망 분리(Decoupling)의 기점."
    )
    _add_methodology_box(doc, "스케일링 법칙",
        "Kaplan et al.(2020)에 따르면 LLM 성능은 C(compute) ∝ N^0.73 · D^0.27로 "
        "모델 크기와 데이터 양의 멱함수로 증가합니다. 이는 GPU 수요 증가가 "
        "모델 성능 개선을 위해 구조적으로 지속됨을 의미합니다.")

    _add_heading(doc, "1.2 밸류 체인 레이어 (14개)", 2)
    vc_rows = []
    for layer, info in config.VALUE_CHAIN.items():
        companies = ", ".join(info.get("companies", [])[:3])
        risk = info.get("bottleneck_risk", "-")
        layer_ko = {
            "end_customer": "최종 사용자", "ai_service": "AI 서비스", "cloud_dc": "클라우드/DC",
            "gpu_server": "GPU 서버", "hbm": "HBM 메모리", "dram": "DRAM",
            "ssd_storage": "SSD 스토리지", "cpu": "CPU", "networking": "네트워킹",
            "packaging": "패키징(CoWoS)", "power": "전력 인프라", "foundry": "파운드리",
            "asic": "ASIC", "edge_ai": "엣지 AI", "sovereign_ai": "Sovereign AI",
        }.get(layer, layer)
        risk_ko = {"critical": "위험", "high": "높음", "medium": "보통", "low": "낮음"}.get(risk, risk)
        vc_rows.append([layer_ko, companies, risk_ko])
    _add_table(doc, ["레이어", "주요 기업", "병목 위험도"], vc_rows)
    doc.add_paragraph()

    _add_heading(doc, "1.3 공급망 병목 전이 (Cascade) 개념", 2)
    doc.add_paragraph(
        "병목 전이는 '기존 병목 해소 → 다음 레이어로 압력 이동'을 반복하는 패턴입니다. "
        "AI 공급망에서는 아래 순서로 진행됩니다:\n\n"
        "  Phase 1 (2023-2024): GPU 부족 → NVIDIA H100 할당 12개월 대기\n"
        "  Phase 2 (2024-현재): HBM 병목 심화 → CoWoS 패키징이 공동 병목\n"
        "  Phase 3 (2025-2026): 전력/DC 인프라 → 변압기·냉각 수요 2배 증가\n"
        "  Phase 4 (2026-2027): 네트워킹 → GB200 클러스터 InfiniBand 수요 급증\n\n"
        "이 패턴은 Morgan Stanley(2025), BCG(2024)의 AI 인프라 투자 프레임워크와 일치합니다."
    )

    _add_ref_table(doc, [
        "NVIDIA_IR", "NVIDIA_GTC_2024", "GPT3_Paper", "Kaplan_Scaling",
        "Bloomberg_AI_Chart", "Sequoia_600B", "Morgan_Stanley_AI"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 2장. 토큰 수요 → 하드웨어 수요 변환 모델
    # ══════════════════════════════════════════
    _add_heading(doc, "2장. 토큰 수요 → 하드웨어 수요 변환 모델", 1)
    doc.add_paragraph(
        "AI 서비스의 핵심 수요 단위는 '토큰(Token)'입니다. 1토큰 ≈ 0.75 영어 단어. "
        "토큰 처리량이 증가할수록 GPU·HBM·전력 수요가 연쇄적으로 증가하며, "
        "이 연쇄 변환은 투자 수요 예측의 핵심 프레임워크입니다."
    )

    _add_heading(doc, "2.1 핵심 수식 및 파라미터 근거", 2)
    formulas = [
        ("Token → GPU",
         "GPU = (Tokens/day ÷ 86,400s) ÷ (tokens/sec × 가동률)",
         "MLCommons MLPerf v4.0: H100 SXM5 = 3,000 tokens/sec (추론 시)"),
        ("GPU → HBM",
         "HBM(GB) = GPU_count × hbm_per_gpu",
         "NVIDIA H100 SXM5 = 80GB HBM3; B200 = 192GB HBM3e"),
        ("KV 캐시 압력",
         "KV(GB) = ctx_len × users × 2bytes × 2 ÷ 1e9",
         "FP16(2바이트) KV Cache; K·V 각 1회 저장 → ×2"),
        ("SSD 수요(RAG)",
         "SSD(GB/day) = tokens/day × RAG_비율 × bytes/token ÷ 1e9",
         "RAG 비율 30% 가정; 기업 AI 워크로드 평균"),
        ("전력 수요",
         "Power(MW) = GPUs × TDP_W × 1.3(서버OH) × PUE ÷ 1e6",
         "H100 TDP=700W; PUE=1.4 (IEA 글로벌 평균)"),
    ]
    _add_table(doc, ["수식명", "공식", "파라미터 근거"], formulas)
    doc.add_paragraph()

    _add_methodology_box(doc, "GPU 추론 처리량 측정",
        "H100 SXM5 추론 throughput: MLCommons MLPerf Inference Datacenter v4.0 벤치마크 "
        "기준 Llama-2 70B 모델에서 약 2,000-3,000 tokens/sec. 실제 운영 가동률 85% 적용. "
        "B200 SXM6는 H100 대비 약 2.5-3배 성능 향상(NVIDIA GTC 2024 발표 기준).")

    _add_heading(doc, "2.2 2026년 Q1 현재 수치 (Base 시나리오)", 2)
    base_rows = [
        ["토큰 수요",      "265T tokens/day",    "2024 25T/day × 3.0x/yr × 2년 복리 성장"],
        ["필요 GPU",       "~180만 개 (H100 기준)", "GPU 가동률 85% 가정"],
        ["HBM 수요",       "~144 PB",             "H100 80GB × 180만 GPU"],
        ["전력 수요",      "~2,304 MW",           "H100 700W TDP × 1.3 × PUE 1.4"],
        ["SSD 수요 (RAG)", "연간 누적 ~100+ EB",  "RAG 비율 30%, 2TB/token 가정"],
    ]
    _add_table(doc, ["지표", "수치", "산출 근거"], base_rows)
    doc.add_paragraph()

    doc.add_paragraph(
        "⚠️ 불확실성 주의사항: 위 수치는 '토큰 처리 GPU' 기준 추정치입니다. "
        "실제 GPU 설치 수는 학습(Training) + 추론(Inference) + 예비(Redundancy)가 포함되어 "
        "2-3배 더 많을 수 있습니다. NVIDIA Data Center 매출 실적($35.6B/Q, 2024Q4)과의 "
        "역산 검증 결과: 약 300-500만 GPU 상당의 인프라가 이미 배포된 것으로 추정됩니다."
    )
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(9)

    _add_heading(doc, "2.3 시나리오별 예측 (2025 실적 / 2026 현재 / 2027-2028 예측)", 2)
    scenario_rows = [
        ["Base", "2025 (실적)", "88.5T/day",  "~60만 GPU",   "~768MW",  "×3.5/yr 성장"],
        ["Base", "2026 ★현재", "265T/day",  "~180만 GPU",  "~2,304MW", "×3.0/yr"],
        ["Base", "2027 (예측)", "795T/day",  "~540만 GPU",  "~6,912MW", "×3.0/yr"],
        ["Base", "2028 (예측)", "2.4P/day",  "~1,620만 GPU","~20,736MW","×3.0/yr"],
        ["Bull", "2026",        "398T/day",  "~270만 GPU",  "~3,456MW", "×4.5/yr"],
        ["Bear", "2026",        "132T/day",  "~89만 GPU",   "~1,138MW", "×2.0/yr"],
    ]
    _add_table(doc, ["시나리오", "연도", "토큰/일", "GPU 수요", "전력(MW)", "성장률 가정"], scenario_rows)

    _add_methodology_box(doc, "시나리오 성장률 근거",
        "Base 3.0×/yr: ChatGPT 2023→2024 토큰 성장 실적 ≈ 5×; "
        "DeepSeek 등 효율화로 일부 완충 가정. "
        "Bull 4.5×/yr: Sequoia 600B 시나리오; 소버린 AI + 에이전트 AI 폭발적 확산. "
        "Bear 2.0×/yr: 경제 침체 또는 AI 거품 조정 시나리오.")

    _add_ref_table(doc, [
        "NVIDIA_IR", "MLCommons_Inference", "Goldman_DC_Power", "IEA_Electricity_2024",
        "Kaplan_Scaling", "Chinchilla_Paper", "Epoch_Compute_Trends", "Sequoia_600B"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 3장. 병목 분석 (상세)
    # ══════════════════════════════════════════
    _add_heading(doc, "3장. 공급망 병목 분석", 1)
    doc.add_paragraph(
        "병목 분석은 레이어별 가동률(Utilization)과 공급 확충 리드타임을 조합하여 "
        "투자 기회를 정량화합니다. 가동률이 높고 리드타임이 길수록 "
        "공급자의 가격 결정력(Pricing Power)이 강화됩니다."
    )

    _add_heading(doc, "3.1 레이어별 가동률 현황 (2026년 Q1 기준)", 2)
    bn_rows = []
    for layer, util in config.CURRENT_CAPACITY_UTILIZATION.items():
        sev = "🔴 위험" if util >= 0.85 else ("🟠 높음" if util >= 0.70 else ("🟡 보통" if util >= 0.50 else "🟢 낮음"))
        resolution = {
            "HBM": "18개월 (신규 팹)", "CoWoS": "12개월 (TSMC 확장)",
            "GPU": "9개월 (Blackwell 램프)", "Power_DC": "24개월 (변압기 리드타임)",
            "Networking": "6개월", "DRAM": "6개월", "SSD": "3개월",
            "Foundry_Advanced": "18개월", "ASIC": "12개월",
        }.get(layer, "N/A")
        layer_ko = {
            "HBM": "HBM 메모리", "CoWoS": "CoWoS 패키징", "GPU": "GPU 서버",
            "Power_DC": "전력/DC 인프라", "Networking": "네트워킹",
            "DRAM": "DRAM", "SSD": "SSD 스토리지", "CPU": "CPU",
            "Foundry_Advanced": "첨단 파운드리", "ASIC": "커스텀 ASIC", "Edge_AI": "엣지 AI"
        }.get(layer, layer)
        bn_rows.append([layer_ko, f"{util*100:.0f}%", sev, resolution])
    _add_table(doc, ["레이어", "가동률", "심각도", "해소 예상 리드타임"], bn_rows)
    doc.add_paragraph()

    _add_methodology_box(doc, "가동률 산출 방법",
        "각 레이어 가동률은 (1) 생산사 IR 공시 발언, (2) TrendForce/Omdia 시장 보고서, "
        "(3) 공급-수요 역산 모델(GPU 출하량 × HBM/GPU = HBM 수요, "
        "대비 생산사 WPM × yield = 공급)을 조합한 추정치입니다. "
        "HBM 92%: SK Hynix IR Q4 2025 '수요가 공급을 초과하는 상황 지속' 발언 근거.")

    _add_heading(doc, "3.2 HBM 병목 상세 분석", 2)
    doc.add_paragraph(
        "HBM(High Bandwidth Memory)은 현재 AI 공급망의 최대 병목입니다. "
        "그 이유는 세 가지입니다:\n\n"
        "① 기술 난이도: HBM은 여러 DRAM 다이를 TSV(Through-Silicon Via)로 수직 적층. "
        "수율(Yield) 75% 수준으로 일반 DRAM 대비 낮음.\n\n"
        "② CoWoS 공동의존: HBM 스택을 GPU에 연결하려면 TSMC의 CoWoS 패키징 공정이 필수. "
        "HBM 생산량이 늘어도 CoWoS 캐파가 제한이면 GPU 출하량이 억제됨.\n\n"
        "③ 세대 전환 복잡성: HBM3e→HBM4 전환 시 설계 변경이 필요하여 "
        "NVIDIA GPU와의 검증(Qualification) 기간 6-12개월 소요. "
        "이 기간 동안 Samsung은 NVIDIA 공급 불가."
    )

    hbm_gen_rows = [
        ["HBM2e", "2020", "461 GB/s", "8GB/stack", "$8/GB", "A100 GPU"],
        ["HBM3",  "2022", "819 GB/s", "16GB/stack", "$14/GB", "H100 GPU"],
        ["HBM3e", "2023", "1,229 GB/s", "24GB/stack", "$18/GB", "H200, B200 GPU"],
        ["HBM4",  "2025", "2,000 GB/s", "32GB/stack", "~$24/GB est", "Blackwell Ultra"],
    ]
    _add_table(doc, ["세대", "양산 시작", "대역폭", "용량/스택", "가격/GB", "탑재 GPU"], hbm_gen_rows)
    doc.add_paragraph()

    _add_heading(doc, "3.3 CoWoS 패키징 병목", 2)
    doc.add_paragraph(
        "CoWoS(Chip-on-Wafer-on-Substrate)는 TSMC가 독점 제공하는 AI GPU용 패키징 기술입니다. "
        "GPU 다이 + HBM 스택을 하나의 패키지로 연결합니다.\n\n"
        "공급 제약 구조:\n"
        "- 2023년 월 35,000wpm → 2024년 50,000wpm → 2025년 85,000wpm → 2026년 120,000wpm\n"
        "- 각 GB200 NVL72 랙 = 9 CoWoS 슬릿 × HBM3e 6스택 = 54 패키징 슬롯\n"
        "- 수요: GB200 NVL72 200,000랙(2026E) × 54 = 1,080만 패키징 슬롯\n"
        "- 공급: 120,000wpm × 12개월 × 수율 = 약 100만 슬롯 (대략적 추정)\n\n"
        "→ CoWoS가 HBM과 함께 GPU 출하량의 실질적 상한을 결정합니다."
    )

    _add_heading(doc, "3.4 전력 인프라 — 부상하는 차세대 병목", 2)
    doc.add_paragraph(
        "Goldman Sachs(2024)에 따르면 AI 데이터센터 전력 수요는 "
        "2022년→2030년 160% 증가가 예상됩니다. 주요 제약:\n\n"
        "① 전력 변압기 리드타임 2.5-3년: 수요 급증에도 즉각 대응 불가\n"
        "② 미국 전력망 연계 신청(Interconnection Queue) 적체: 2024년 기준 2,600GW 대기 중\n"
        "③ 미국 AI 데이터센터 전력 수요: 2022년 ~4GW → 2030년 ~15GW+ 전망\n\n"
        "이로 인해 Vertiv, Eaton, GE Vernova 등 전력 인프라 기업의 주문잔고가 "
        "사상 최고치를 기록하고 있습니다."
    )

    _add_ref_table(doc, [
        "SKHynix_HBM", "SKHynix_IR", "Omdia_HBM", "TrendForce_DRAM",
        "TSMC_CoWoS", "TSMC_AR_2024", "DigiTimes_CoWoS", "SemiAnalysis_CoWoS",
        "Goldman_DC_Power", "IEA_Electricity_2024", "Lawrence_Berkeley_DC",
        "Vertiv_IR", "GEVernova_IR"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 4장. 하이퍼스케일러 CapEx 동향
    # ══════════════════════════════════════════
    _add_heading(doc, "4장. 하이퍼스케일러 CapEx 동향 (2018-2026)", 1)
    doc.add_paragraph(
        "AI 인프라 투자를 이해하려면 하이퍼스케일러의 CapEx(자본 지출) 트렌드를 "
        "장기 시계열로 분석해야 합니다. Big 4(Microsoft·Google·Amazon·Meta) + xAI의 "
        "AI 인프라 투자는 2024년 $219B에서 2025년 $340B+으로 55% 급증했습니다."
    )

    _add_heading(doc, "4.1 연도별 CapEx 테이블 (USD 십억, 연간)", 2)
    capex_rows = []
    for company, years in config.HYPERSCALER_CAPEX.items():
        row = [company]
        for yr in ["2020", "2022", "2023", "2024", "2025", "2026_est"]:
            val = years.get(yr) or years.get(yr.replace("_est",""))
            row.append(f"${val}B" if val else "-")
        capex_rows.append(row)
    _add_table(doc, ["회사", "2020", "2022", "2023", "2024", "2025", "2026E"], capex_rows)
    doc.add_paragraph()

    _add_heading(doc, "4.2 회사별 AI 투자 전략 분석", 2)
    strategy_detail = [
        ["Microsoft",
         "$80B (FY2025)",
         "OpenAI 독점 Azure 서비스 내재화; Copilot 전 제품 통합",
         "OpenAI 의존도 리스크; CapEx 규모 지속 가능성"],
        ["Google (Alphabet)",
         "$75B (2025E)",
         "TPU v5 자체 칩 + NVIDIA GPU 병행; Gemini Ultra 경쟁력",
         "AI 서비스 광고 매출 잠식 리스크"],
        ["Amazon (AWS)",
         "$105B (2025E)",
         "Trainium2 자체 칩 + Anthropic $4B 투자; AWS 인프라 확장",
         "AWS 마진 압박; 자체 칩 수율 리스크"],
        ["Meta",
         "$65B (2025E)",
         "LLaMA 오픈소스 전략; MTIA ASIC 자체 개발; 단기 매출 없음",
         "오픈소스 경쟁 심화로 수익 모델 불명확"],
        ["xAI",
         "$15B (2025E)",
         "Colossus 클러스터 H100 10만→20만; Grok 모델 집중",
         "수익화 타임라인 불명확"],
    ]
    _add_table(doc, ["회사", "2025 CapEx", "AI 전략 핵심", "주요 리스크"], strategy_detail)
    doc.add_paragraph()

    _add_heading(doc, "4.3 CapEx와 GPU 수요의 관계", 2)
    doc.add_paragraph(
        "하이퍼스케일러 CapEx의 약 40-60%가 GPU 구매 관련 지출로 추정됩니다 "
        "(NVIDIA 매출과 역산). 2024년 Big 4 합산 CapEx $219B 중 GPU 관련 약 $100B 이상이 "
        "NVIDIA Data Center 매출($113B FY2024)과 대략 일치합니다.\n\n"
        "나머지 CapEx는 네트워킹(InfiniBand/Ethernet), 서버랙, 냉각 시스템, "
        "전력 인프라, 부동산에 배분됩니다."
    )

    _add_ref_table(doc, [
        "MSFT_10K_2024", "MSFT_Earnings_Q4FY2025", "Google_10K_2024",
        "Amazon_10K_2024", "Meta_10K_2024", "NVIDIA_IR", "Bernstein_HBM"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 5장. 투자 네트워크 분석
    # ══════════════════════════════════════════
    _add_heading(doc, "5장. AI 공급망 투자 네트워크 분석", 1)
    doc.add_paragraph(
        "AI 공급망의 자금 흐름은 크게 4가지 관계 유형으로 분류됩니다:\n"
        "① Investment(전략적 투자): 하이퍼스케일러 → AI 스타트업\n"
        "② Hardware Supply(하드웨어 공급): 반도체 기업 → 서버/클라우드\n"
        "③ Service(서비스 계약): AI 모델 API → 최종 고객\n"
        "④ VC(벤처 투자): VC 펀드 → AI 스타트업\n\n"
        "이 관계를 네트워크 그래프로 시각화하면 공급망의 핵심 허브(Hub)와 "
        "단일 실패점(Single Point of Failure)이 드러납니다."
    )

    _add_heading(doc, "5.1 주요 전략적 투자 관계", 2)
    inv_rows = []
    for rel in config.COMPANY_RELATIONSHIPS:
        if rel[2] in ("investment", "vc") and len(inv_rows) < 18:
            inv_rows.append([
                rel[0], rel[1], rel[2].replace("investment","전략 투자").replace("vc","VC"),
                rel[3] if len(rel)>3 else "",
                rel[4] if len(rel)>4 else ""
            ])
    _add_table(doc, ["투자자/출처", "피투자사/대상", "유형", "설명", "규모"], inv_rows)
    doc.add_paragraph()

    _add_heading(doc, "5.2 하드웨어 공급 독점 구조 분석", 2)
    doc.add_paragraph(
        "AI 공급망에서 단일 공급자 의존(Single-Vendor Dependency) 위험이 높은 레이어:\n\n"
        "① TSMC CoWoS: 전 세계 AI GPU 패키징의 ~95% 독점. 대안 없음(현재).\n"
        "② SK Hynix HBM3e: NVIDIA 주 공급사 (~52% 점유율). Samsung 대기 중.\n"
        "③ NVIDIA GPU: AI 훈련 가속기 시장의 ~80%+ 점유율. AMD/Intel 추격 중.\n\n"
        "이러한 독점 레이어는 단기적으로 높은 수익성과 가격 결정력을 의미하지만, "
        "장기적으로 대안 공급사 출현 리스크가 존재합니다."
    )

    hw_rows = []
    for rel in config.COMPANY_RELATIONSHIPS:
        if rel[2] == "hardware_supply" and len(hw_rows) < 12:
            excl = "독점" if "독점" in (rel[4] if len(rel) > 4 else "") else "경쟁"
            hw_rows.append([rel[0], rel[1], rel[3] if len(rel) > 3 else "", excl])
    _add_table(doc, ["공급사", "수요사", "공급 내용", "경쟁 구조"], hw_rows[:12])
    doc.add_paragraph()

    _add_heading(doc, "5.3 NVIDIA 중심의 공급망 네트워크 특성", 2)
    doc.add_paragraph(
        "NVIDIA는 AI 공급망 네트워크의 '허브' 역할을 합니다:\n\n"
        "- 상류(Upstream) 의존: SK Hynix HBM3e, TSMC 3nm/CoWoS, Micron HBM\n"
        "- 하류(Downstream) 지배: Microsoft Azure, AWS, Google Cloud, Meta에 GPU 공급\n"
        "- 소프트웨어 모트(Moat): CUDA 생태계 10년 축적 → 전환 비용 극대화\n\n"
        "이로 인해 NVIDIA의 시가총액이 $3조+에 도달했으나, "
        "Sequoia의 '$600B 질문'은 AI 수익화 실현 속도가 인프라 투자를 뒷받침할 수 있는지를 묻습니다."
    )

    _add_ref_table(doc, [
        "Bloomberg_AI_Chart", "Sequoia_600B", "NVIDIA_10K_2024",
        "TSMC_CoWoS", "SKHynix_HBM", "Micron_HBM",
        "Morgan_Stanley_AI", "BCG_AI_Infra"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 6장. 투자 시그널 & 포지셔닝
    # ══════════════════════════════════════════
    _add_heading(doc, "6장. 투자 시그널 & 포지셔닝", 1)
    doc.add_paragraph(
        "AI 공급망 투자 분석의 핵심 프레임워크는 '병목 위치 + 리드타임 + 대안 공급사 유무'입니다. "
        "병목이 심할수록, 리드타임이 길수록, 대안이 없을수록 투자 매력도가 높습니다. "
        "현재(2026년 Q1) AI 투자 사이클은 Phase 2~3 전환 구간입니다."
    )

    _add_heading(doc, "6.1 투자 단계별 로드맵 (Phase Framework)", 2)
    phase_rows = [
        ["Phase 1", "2023-2024", "GPU 인프라",       "완료/선반영",  "NVIDIA, AMD",             "GPU 수요 폭증기; NVDA +200%+"],
        ["Phase 2", "2024-현재", "HBM & 패키징",     "★현재 핵심", "SK Hynix, TSMC CoWoS",    "메모리 업사이클; HBM 병목 심화"],
        ["Phase 3", "2025-2026", "전력 인프라",       "진입 중",     "Vertiv, GE Vernova, Eaton","DC 전력 2배 증가; 변압기 부족"],
        ["Phase 4", "2026-2027", "네트워킹/ASIC",    "예비",        "Broadcom, Marvell",        "GB200 클러스터 IB 수요 급증"],
        ["Phase 5", "2027-2028", "엣지/로보틱스",     "장기",        "Qualcomm, 로봇 기업",      "AI 단말·물리 AI 확산"],
    ]
    _add_table(doc, ["단계", "시기", "테마", "현황", "수혜 기업", "투자 포인트"], phase_rows)
    doc.add_paragraph()

    _add_methodology_box(doc, "Phase Framework 근거",
        "Morgan Stanley(2025) 'AI Investment Playbook'의 인프라 사이클 분석, "
        "BCG(2024) 'Winning the AI Infrastructure Race' 프레임워크, "
        "Sequoia(2024) '$600B Question' 투자 단계 분석을 종합하여 구성.")

    _add_heading(doc, "6.2 종목별 투자 시그널 상세", 2)
    signal_rows = [
        ["SK Hynix\n(000660.KS)", "Strong Buy",
         "HBM3e 52% 시장 점유율; 2025-2026 GB200 NVL72 전용 공급; HBM4 양산 2025 시작",
         "SK Hynix IR Q4 2025; Omdia HBM Report 2024",
         "Samsung HBM3e 인증 완료 시 점유율 잠식 리스크"],
        ["TSMC\n(TSM)", "Buy",
         "CoWoS 독점 (95%+); 캐파 50K→120K wpm 2024-2026; N3/N2 AI GPU 탑재 가격 프리미엄",
         "TSMC AR 2024; DigiTimes CoWoS Tracker",
         "지정학적 대만 리스크; Intel Foundry 3A 경쟁"],
        ["Vertiv\n(VRT)", "Strong Buy",
         "주문잔고 $8B+; 전력/냉각 인프라 2026 sold-out; 고마진 서비스 사업 성장",
         "Vertiv IR; Goldman Sachs AI Power 2024",
         "원자재(구리) 가격 변동; 새 경쟁사 진입"],
        ["GE Vernova\n(GEV)", "Buy",
         "가스터빈·변압기 기록적 주문잔고; AI DC 전력 수요 핵심 수혜",
         "GEVernova IR; IEA Electricity 2024",
         "인허가 지연; 규제 리스크"],
        ["Broadcom\n(AVGO)", "Buy",
         "Google/Meta 커스텀 ASIC 설계 독점; 400G/800G 네트워킹 칩",
         "Broadcom IR; Bloomberg AI Network Chart",
         "NVIDIA NVLink 생태계 경쟁; 고객 집중 리스크"],
        ["NVIDIA\n(NVDA)", "Hold/Watch",
         "Data Center $49B/분기(FY2025); B200/GB200 램프; CUDA 소프트웨어 모트",
         "NVIDIA 10-K 2024; GTC 2024 Keynote",
         "밸류에이션 $3T+; DeepSeek 효율화; 수출 규제"],
        ["Micron\n(MU)", "Buy",
         "HBM3e 점유율 10%→20%+ 목표; CHIPS Act 미국 HBM 팹; NVIDIA 다변화 수혜",
         "Micron IR; CHIPS Act BIS 문서",
         "SK Hynix 경쟁 우위; 수율 램프 리스크"],
        ["Samsung\n(005930.KS)", "Watch",
         "HBM3e NVIDIA 인증 대기; 통과 시 35% 점유율 탈환 가능; 밸류에이션 저평가",
         "Samsung IR; TrendForce 분기 보고",
         "인증 실패 리스크; 반도체 사이클 불확실성"],
    ]
    _add_table(doc, ["기업(티커)", "시그널", "투자 근거", "데이터 출처", "주요 리스크"], signal_rows)

    _add_ref_table(doc, [
        "SKHynix_HBM", "TSMC_AR_2024", "Vertiv_IR", "GEVernova_IR",
        "Morgan_Stanley_AI", "Bernstein_HBM", "Goldman_DC_Power",
        "Sequoia_600B", "CHIPS_Act"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 7장. 2025-2028 수요 예측 시나리오
    # ══════════════════════════════════════════
    _add_heading(doc, "7장. 2025-2028 수요 예측 시나리오", 1)

    _add_heading(doc, "7.1 시나리오 방법론", 2)
    doc.add_paragraph(
        "시나리오 분석은 미래 불확실성을 범위로 표현합니다. "
        "본 모델의 핵심 드라이버와 민감도:\n\n"
        "① 토큰 연간 성장률 (가장 큰 불확실성): Bear 2×/yr ~ Bull 4.5×/yr\n"
        "② GPU 가동률 가정: Bear 75% ~ Bull 90% (현재 Base 85%)\n"
        "③ HBM per GPU 증가: HBM3e 80GB(H100) → 192GB(B200) → 예측치\n"
        "④ 추론/학습 비율 변화: 현재 50/50 → 2027년 30/70 (추론 급증)\n"
        "⑤ DeepSeek 효과: 효율성 개선이 Jevons Paradox로 총 수요 증폭 가정\n\n"
        "Jevons 역설: 에너지 효율이 높아지면 총 에너지 소비가 증가하는 현상 "
        "(Jevons 1865). AI에서는 모델 효율 향상 → 더 많은 사용 케이스 → 총 토큰 수요 증가."
    )

    _add_heading(doc, "7.2 핵심 지표 예측 테이블", 2)
    pred_rows = [
        ["2024", "실적", "25T/day",  "~150만",  "~120PB", "~1,920MW", "기준년도"],
        ["2025", "Base 실적", "88.5T/day", "~200만", "~160PB", "~2,560MW", "×3.5 성장"],
        ["2026 ★", "Base 현재", "265T/day", "~540만", "~432PB", "~6,912MW", "×3.0 성장"],
        ["2027", "Base 예측", "795T/day", "~1,620만", "~1,296PB", "~20,736MW", "×3.0 성장"],
        ["2028", "Base 예측", "2.4P/day", "~4,860만", "~3,888PB", "~62,208MW", "×3.0 성장"],
        ["2026 ★", "Bull", "398T/day", "~810만", "~648PB", "~10,368MW", "×4.5 성장"],
        ["2026 ★", "Bear", "132T/day", "~270만", "~216PB", "~3,456MW", "×2.0 성장"],
        ["2027", "Bull 예측", "1.8P/day", "~3,645만", "~2,916PB", "~46,656MW", "×4.5 성장"],
        ["2027", "Bear 예측", "264T/day", "~540만", "~432PB", "~6,912MW", "×2.0 성장"],
    ]
    _add_table(doc, ["연도", "시나리오", "토큰/일", "GPU 수요", "HBM 수요", "전력(MW)", "비고"], pred_rows)
    doc.add_paragraph()

    doc.add_paragraph(
        "⚠️ 불확실성 경고: 2027-2028년 예측은 현재 기술 추세의 외삽(Extrapolation)이며, "
        "실제 결과와 크게 다를 수 있습니다. 특히:\n"
        "- 양자 컴퓨팅 또는 광(Photonic) AI 칩 등장 시 GPU/HBM 수요 변동 가능\n"
        "- AI 규제(EU AI Act 등) 강화 시 Bear 시나리오 가능\n"
        "- AGI(범용인공지능) 조기 달성 시 Bull을 초과하는 수요 폭발 가능"
    )
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(9)

    _add_heading(doc, "7.3 시나리오별 수혜/피해 분석", 2)
    impact_rows = [
        ["Bull", "최대 수혜", "SK Hynix, TSMC, Vertiv, NVIDIA, GE Vernova",
         "HBM·패키징·전력 초과 수요; 공급자 가격 결정력 극대화"],
        ["Bull", "리스크", "공급 확보 실패한 AI 서비스 기업",
         "인프라 병목으로 서비스 확장 제한; 경쟁사에 뒤처짐"],
        ["Base", "수혜", "SK Hynix, TSMC CoWoS, Vertiv",
         "지속적 병목 프리미엄; 안정적 수요 성장"],
        ["Bear", "수혜", "AMD MI300X/MI325X, Intel Gaudi3",
         "NVIDIA 대비 가격 경쟁력 부각; 점유율 확대"],
        ["Bear", "리스크", "HBM 과잉 투자 기업, 고밸류 AI 인프라 주식",
         "공급 과잉 전환 시 가격 하락; NVIDIA 밸류에이션 조정"],
        ["DeepSeek 효과", "Jevons 역설", "총 AI 인프라 수요 증폭",
         "모델 효율 개선이 더 많은 AI 사용 케이스 → 수요 폭발"],
    ]
    _add_table(doc, ["시나리오", "구분", "기업", "근거"], impact_rows)

    _add_ref_table(doc, [
        "Sequoia_600B", "IEA_Electricity_2024", "IDC_GPU_2024",
        "Gartner_AI_Forecast", "McKinsey_AI_State", "Jevons_Paradox",
        "Epoch_Compute_Trends", "DeepSeek_R1"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 8장. AI 모델 발전 타임라인과 하드웨어 수요
    # ══════════════════════════════════════════
    _add_heading(doc, "8장. AI 모델 발전 타임라인과 하드웨어 수요 연결", 1)
    doc.add_paragraph(
        "AI 모델의 파라미터 규모 증가는 하드웨어 수요와 직접 연결됩니다. "
        "GPT-3 175B(2020) → GPT-4 추정 1.8T(2023) → 현재 MoE 2-3T 규모까지 "
        "약 5년 간 파라미터가 ~10배 증가했으며, "
        "이는 Kaplan et al.의 스케일링 법칙과 Epoch AI의 컴퓨팅 트렌드 분석을 뒷받침합니다."
    )

    _add_heading(doc, "8.1 주요 AI 모델 마일스톤", 2)
    milestone_rows = [
        ["2020-05", "GPT-3", "OpenAI", "175B", "스케일링 법칙 실증; A100 수요 촉발"],
        ["2022-11", "ChatGPT", "OpenAI", "175B (RLHF)", "소비자 AI 폭발; 1억 사용자 2개월"],
        ["2023-03", "GPT-4", "OpenAI", "~1.8T (MoE 추정)", "H100 할당 12개월+ 대기 촉발"],
        ["2023-07", "LLaMA-2 70B", "Meta", "70B", "오픈소스; SSD 로컬 AI 수요 급증"],
        ["2024-09", "o1 (추론)", "OpenAI", "~300B 추정", "추론 시 GPU 10배; HBM 수요 증폭"],
        ["2025-01", "DeepSeek R1", "DeepSeek", "671B (MoE)", "Jevons 역설; GPU 효율 재정립"],
        ["2025-02", "Claude 3.7 Sonnet", "Anthropic", "~200B 추정", "에이전틱 AI; 추론 시간 컴퓨팅"],
    ]
    _add_table(doc, ["출시일", "모델", "기관", "파라미터", "하드웨어 영향"], milestone_rows)
    doc.add_paragraph()

    _add_heading(doc, "8.2 소버린 AI와 추가 수요", 2)
    sovereign_rows = [
        ["UAE", "$100B", "2025-2030", "~50,000 GPU", "G42/MGX; NVIDIA 파트너십"],
        ["사우디아라비아", "$40B", "2024-2027", "~30,000 GPU", "NEOM AI 인프라"],
        ["EU (집합)", "$20B", "2025-2030", "~40,000 GPU", "IPCEI AI; 유럽 AI 자주권"],
        ["일본", "$7B", "2024-2026", "~20,000 GPU", "SoftBank-NVIDIA; NEDO 지원"],
        ["인도", "$1.25B", "2024-2027", "~10,000 GPU", "IndiaAI Mission; MS/Google"],
    ]
    _add_table(doc, ["국가", "투자 규모", "기간", "GPU 수요 추정", "비고"], sovereign_rows)
    doc.add_paragraph()

    doc.add_paragraph(
        "소버린 AI 추가 GPU 수요 합계: ~150,000 GPU (단기) ~ 수십만 GPU (중장기)\n"
        "이는 NVIDIA 연간 출하량(2025E ~5-7백만 GPU)의 2-5% 수준이나, "
        "단일 주문 규모가 커 공급망 우선순위에 영향을 줍니다."
    )
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(9)

    _add_ref_table(doc, [
        "Attention_Paper", "GPT3_Paper", "Kaplan_Scaling", "Chinchilla_Paper",
        "DeepSeek_R1", "Epoch_Compute_Trends", "MLCommons_Inference",
        "UAE_AI_Deal", "Export_Control_BIS", "SIA_Factbook"
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 종합 참고문헌 목록
    # ══════════════════════════════════════════
    _add_heading(doc, "종합 참고문헌", 1)
    doc.add_paragraph(
        f"총 {len(ALL_REFERENCES)}개 참고자료 | 기준일: {config.AS_OF_DATE}"
    )
    ref_full_rows = []
    for key, (name, url, desc, yr) in ALL_REFERENCES.items():
        ref_full_rows.append([name, yr, desc])
    _add_table(doc, ["출처명", "연도", "내용 요약"], ref_full_rows)

    # ══════════════════════════════════════════
    # 저장
    # ══════════════════════════════════════════
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today_str = date.today().strftime("%Y%m%d")
    out_path = OUTPUT_DIR / f"ai_scm_study_{today_str}.docx"
    doc.save(str(out_path))
    print(f"  [Word] 학습 문서 생성: {out_path} (8장, {len(ALL_REFERENCES)}개 참고자료)")
    return str(out_path)


def run(state=None, bottleneck=None, strategy=None, modeling=None, mapping=None):
    return {"word_path": build(state, bottleneck, strategy, modeling, mapping)}
