"""
lecture/agents/word_builder.py — AI 뉴스 브리핑 Word 문서 생성

python-docx 기반 전문 브리핑 문서 생성:
  - 헤더/날짜/소스 정보
  - 뉴스 요약 표
  - 뉴스별 상세 분석
  - 트렌드 분석
  - 강의 추천 업데이트 아이디어

저장 위치: output/daily/news_briefing_YYYYMMDD.docx
"""

import pathlib
import sys
from datetime import datetime

sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))
from config import OUTPUT_DIR

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


DAILY_DIR = OUTPUT_DIR / "daily"

# 색상 정의 (RGBColor)
C_NAVY  = RGBColor(0x0A, 0x16, 0x28)
C_BLUE  = RGBColor(0x1D, 0x4E, 0xD8)
C_MUTED = RGBColor(0x64, 0x74, 0x8B)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BORDER= RGBColor(0xE2, 0xE8, 0xF0)


def _set_cell_bg(cell, hex_color: str) -> None:
    """표 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_para_spacing(para, before: int = 0, after: int = 0) -> None:
    """단락 간격 설정 (pt)."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"), str(after * 20))
    pPr.append(spacing)


def _relevance_color(relevance: str) -> str:
    return {"높음": "1D4ED8", "보통": "0D6B4F", "낮음": "94A3B8"}.get(relevance, "94A3B8")


def build_daily_briefing_doc(briefing_data: dict, date_str: str | None = None) -> str:
    """
    뉴스 브리핑 Word 문서 생성.

    Parameters:
      briefing_data: news_agent.run_news_agent() 반환값
      date_str: 날짜 문자열 (None이면 오늘)

    Returns:
      저장된 .docx 파일 경로
    """
    if not DOCX_AVAILABLE:
        print("  [오류] python-docx 미설치 — pip3 install python-docx")
        return ""

    DAILY_DIR.mkdir(parents=True, exist_ok=True)
    today = date_str or datetime.now().strftime("%Y-%m-%d")
    filename = f"news_briefing_{today.replace('-', '')}.docx"
    out_path = DAILY_DIR / filename

    analysis = briefing_data.get("analysis", {})
    new_items = briefing_data.get("new_items", [])
    top_news = analysis.get("top_news", [])

    doc = Document()

    # ── 페이지 여백 설정 ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 제목 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(title, before=0, after=4)
    run = title.add_run("AI 워크플로우 뉴스 브리핑")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = C_NAVY

    # 날짜 / 출처
    sub = doc.add_paragraph()
    _set_para_spacing(sub, before=0, after=8)
    run2 = sub.add_run(f"{today}  ·  수집 뉴스 {len(new_items)}건  ·  AI Workflow Intelligence Report")
    run2.font.size = Pt(10)
    run2.font.color.rgb = C_MUTED

    # 구분선
    doc.add_paragraph("─" * 80)

    # ── 주요 통계 ──
    stats_para = doc.add_paragraph()
    _set_para_spacing(stats_para, before=4, after=4)
    stats_run = stats_para.add_run(
        f"신규 수집: {len(new_items)}건   |   "
        f"긴급도: {analysis.get('urgency', '보통')}   |   "
        f"강의 업데이트 필요: {'예' if briefing_data.get('has_significant_updates') else '아니오'}"
    )
    stats_run.font.size = Pt(11)
    stats_run.font.bold = True
    stats_run.font.color.rgb = C_BLUE

    # ── 트렌드 요약 ──
    if analysis.get("trend_summary"):
        doc.add_paragraph()
        trend_hd = doc.add_paragraph()
        _set_para_spacing(trend_hd, before=6, after=2)
        th_run = trend_hd.add_run("📊  이번 주 AI 트렌드 요약")
        th_run.font.size = Pt(13)
        th_run.font.bold = True
        th_run.font.color.rgb = C_NAVY

        trend_body = doc.add_paragraph(analysis["trend_summary"])
        trend_body.runs[0].font.size = Pt(11)
        trend_body.runs[0].font.color.rgb = C_MUTED
        _set_para_spacing(trend_body, before=0, after=6)

    # ── 주요 뉴스 요약 표 ──
    if top_news:
        doc.add_paragraph()
        table_hd = doc.add_paragraph()
        _set_para_spacing(table_hd, before=6, after=4)
        th_run2 = table_hd.add_run("📰  주요 뉴스 요약")
        th_run2.font.size = Pt(13)
        th_run2.font.bold = True
        th_run2.font.color.rgb = C_NAVY

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # 헤더 행
        hdr_cells = table.rows[0].cells
        headers = ["순위", "제목", "출처", "강의 관련성"]
        widths = [Cm(1.5), Cm(10.0), Cm(3.0), Cm(2.5)]
        for cell, text, width in zip(hdr_cells, headers, widths):
            _set_cell_bg(cell, "0A1628")
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = C_WHITE
            cell.width = width

        # 뉴스 행
        for news in top_news[:10]:
            row = table.add_row()
            cells = row.cells
            row_data = [
                str(news.get("rank", "")),
                news.get("title", "")[:80],
                news.get("source", ""),
                news.get("lecture_relevance", ""),
            ]
            for i, (cell, text) in enumerate(zip(cells, row_data)):
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.font.size = Pt(9)
                if i == 3:  # 관련성 컬럼 색상
                    run.font.color.rgb = RGBColor(
                        *tuple(int(_relevance_color(text)[j:j+2], 16) for j in (0, 2, 4))
                    )
                    run.font.bold = True

    # ── 뉴스별 상세 분석 ──
    if top_news:
        doc.add_paragraph()
        doc.add_page_break()

        detail_hd = doc.add_paragraph()
        _set_para_spacing(detail_hd, before=0, after=8)
        dh_run = detail_hd.add_run("📋  뉴스별 상세 분석")
        dh_run.font.size = Pt(16)
        dh_run.font.bold = True
        dh_run.font.color.rgb = C_NAVY

        for news in top_news[:8]:
            # 뉴스 제목
            n_title = doc.add_paragraph()
            _set_para_spacing(n_title, before=10, after=2)
            nt_run = n_title.add_run(f"[{news.get('rank', '')}] {news.get('title', '')}")
            nt_run.font.size = Pt(12)
            nt_run.font.bold = True
            nt_run.font.color.rgb = C_NAVY

            # 메타
            meta = doc.add_paragraph()
            _set_para_spacing(meta, before=0, after=2)
            meta.add_run(f"출처: {news.get('source', '')}  |  강의 관련성: {news.get('lecture_relevance', '')}").font.size = Pt(9)
            meta.runs[0].font.color.rgb = C_MUTED

            # 요약
            if news.get("summary"):
                summary_p = doc.add_paragraph(news["summary"])
                _set_para_spacing(summary_p, before=2, after=2)
                summary_p.runs[0].font.size = Pt(10)

            # 강의 활용 포인트
            if news.get("lecture_note"):
                note_p = doc.add_paragraph()
                _set_para_spacing(note_p, before=2, after=2)
                note_p.add_run("강의 활용: ").font.bold = True
                note_p.runs[0].font.size = Pt(10)
                note_p.runs[0].font.color.rgb = C_BLUE
                note_p.add_run(news["lecture_note"]).font.size = Pt(10)

            # 슬라이드 아이디어
            if news.get("slide_idea"):
                idea_p = doc.add_paragraph()
                _set_para_spacing(idea_p, before=2, after=4)
                idea_p.add_run("💡 슬라이드 아이디어: ").font.bold = True
                idea_p.runs[0].font.size = Pt(10)
                idea_p.runs[0].font.color.rgb = RGBColor(0x0D, 0x6B, 0x4F)
                idea_p.add_run(news["slide_idea"]).font.size = Pt(10)

    # ── 주요 개발 사항 ──
    key_devs = analysis.get("key_developments", [])
    if key_devs:
        doc.add_paragraph()
        kd_hd = doc.add_paragraph()
        _set_para_spacing(kd_hd, before=8, after=4)
        kd_hd.add_run("🔑  주요 개발 사항").font.bold = True
        kd_hd.runs[0].font.size = Pt(13)
        kd_hd.runs[0].font.color.rgb = C_NAVY
        for dev in key_devs:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(dev).font.size = Pt(10)
            _set_para_spacing(bullet, before=1, after=1)

    # ── 강의 추천 업데이트 ──
    recs = analysis.get("recommended_additions", [])
    if recs:
        doc.add_paragraph()
        rec_hd = doc.add_paragraph()
        _set_para_spacing(rec_hd, before=8, after=4)
        rec_hd.add_run("📌  강의 자료 업데이트 추천").font.bold = True
        rec_hd.runs[0].font.size = Pt(13)
        rec_hd.runs[0].font.color.rgb = C_NAVY
        for rec in recs:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(rec).font.size = Pt(10)
            _set_para_spacing(bullet, before=1, after=1)

    # ── 푸터 ──
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    footer_p = doc.add_paragraph()
    footer_p.add_run(
        f"생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        "Powered by Claude Code + Anthropic  |  AI Workflow Intelligence Report"
    ).font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = C_MUTED

    doc.save(str(out_path))
    print(f"  ✅ Word 브리핑 저장 → {out_path}")
    return str(out_path)


def build_all_demos() -> list[str]:
    """데모용 — 샘플 데이터로 Word 문서 생성 테스트."""
    sample = {
        "new_items": [{"id": "abc", "title": "OpenAI releases GPT-5", "source": "TechCrunch"}],
        "analysis": {
            "top_news": [
                {
                    "rank": 1,
                    "title": "OpenAI releases GPT-5 with 10x reasoning improvements",
                    "source": "TechCrunch",
                    "summary": "OpenAI가 GPT-5를 공개했습니다. 이전 버전 대비 추론 능력이 10배 향상되었으며 멀티모달 처리 성능이 대폭 개선되었습니다.",
                    "lecture_relevance": "높음",
                    "lecture_note": "AI 진화 챕터에서 최신 모델 소개 시 활용 가능",
                    "slide_idea": "GPT-5 vs 이전 버전 성능 비교 슬라이드",
                }
            ],
            "trend_summary": "이번 주 AI 업계는 멀티모달 AI와 에이전트 기술의 발전이 두드러졌습니다.",
            "key_developments": ["GPT-5 출시", "Claude 3.8 베타 테스트", "Google Gemini Ultra 업데이트"],
            "recommended_additions": ["GPT-5 성능 데이터 추가", "최신 비용 비교표 업데이트"],
            "urgency": "높음",
        },
        "has_significant_updates": True,
        "total_collected": 15,
    }
    path = build_daily_briefing_doc(sample)
    return [path] if path else []


if __name__ == "__main__":
    print("📝 샘플 Word 브리핑 생성 중...")
    paths = build_all_demos()
    for p in paths:
        print(f"  → {p}")
